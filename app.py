# claudiafito_v2 — Atendimento + Binaural (como no app antigo)
# Single-file Streamlit app (compatível com Python 3.9+)kk

import os
import json
import base64
import io
import wave
from datetime import date, timedelta
from typing import Dict, List, Tuple, Any, Optional

import math

import numpy as np
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

# --- Optional dependency: reportlab (PDF)
try:
    import reportlab  # type: ignore
    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False


# -----------------------------
# Config
# -----------------------------
st.set_page_config(page_title="claudiafito_v2", layout="wide")

def _get_env_or_secret(key: str) -> Optional[str]:
    try:
        if hasattr(st, "secrets") and key in st.secrets:
            return st.secrets[key]
    except Exception:
        pass
    return os.getenv(key)

DATABASE_URL = _get_env_or_secret("DATABASE_URL")
SUPABASE_URL = _get_env_or_secret("SUPABASE_URL")
SUPABASE_KEY = _get_env_or_secret("SUPABASE_SERVICE_ROLE_KEY") or _get_env_or_secret("SUPABASE_KEY")

BACKEND = "postgres" if DATABASE_URL else ("supabase" if (SUPABASE_URL and SUPABASE_KEY) else "none")
if BACKEND == "none":
    st.error("Defina DATABASE_URL **OU** SUPABASE_URL + SUPABASE_KEY (preferencialmente SUPABASE_SERVICE_ROLE_KEY).")
    st.stop()

# -----------------------------
# DB helpers
# -----------------------------
if BACKEND == "postgres":
    import psycopg2
    import psycopg2.extras

    def get_conn():
        return psycopg2.connect(DATABASE_URL)

    def qall(sql: str, params=None):
        with get_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(sql, params)
                return cur.fetchall()

    def qone(sql: str, params=None):
        rows = qall(sql, params)
        return rows[0] if rows else None

    def qexec(sql: str, params=None):
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(sql, params)
            conn.commit()

else:
    from supabase import create_client, Client  # pip install supabase

    sb: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

    def sb_select(
        table: str,
        columns: str = "*",
        eq: Optional[Dict[str, Any]] = None,
        order: Optional[Tuple[str, bool]] = None,
        limit: Optional[int] = None,
    ):
        q = sb.table(table).select(columns)
        if eq:
            for k, v in eq.items():
                q = q.eq(k, v)
        if order:
            col, asc = order
            q = q.order(col, desc=not asc)
        if limit is not None:
            q = q.limit(limit)
        res = q.execute()
        return res.data or []

    def sb_insert(table: str, payload: Dict[str, Any]):
        res = sb.table(table).insert(payload).execute()
        data = res.data or []
        return data[0] if data else None


# -----------------------------
# Domain model (anamnese simples)
# -----------------------------
DOMAINS = ["sono", "ansiedade", "humor_baixo", "exaustao", "pertencimento", "tensao", "ruminacao"]
# -----------------------------
# Escala 0–4 (para todos os sliders da anamnese)
# 0 = sem sintoma / ok (bom)  |  4 = muito intenso / sempre (ruim)
# -----------------------------
SCALE_0_4_HELP = (
    "📏 **Escala 0–4** usada nas perguntas: "
    "**0 = nada/sem sintoma (bom)** · 1 = leve · 2 = moderado · 3 = forte · **4 = muito forte/sempre (ruim)**. "
    "Ex.: *Dificuldade para pegar no sono* → 0 = dorme bem, 4 = quase sempre tem dificuldade."
)




# -----------------------------
# V3 — Atendimento (Mapa 0–10) + Neuroplasticidade
# Observação: mantemos DOMAINS/protocolos existentes para NÃO mudar a base de dados.
# O mapa 0–10 é convertido em scores (0–100%) por domínio, usados na seleção de protocolos e sessões.
# -----------------------------
SCALE_0_10_HELP = (
    "📏 **Escala 0–10** (mapa rápido): **0 = nada/ok (bom)** · 5 = moderado · **10 = muito intenso/sempre (ruim)**."
)

V3_SYMPTOMS = [
    # Corpo
    {"id": "sym_dor", "label": "Dor no corpo (geral)", "group": "corpo"},
    {"id": "sym_rigidez", "label": "Rigidez / travamento (articulações / músculos)", "group": "corpo"},
    {"id": "sym_sono", "label": "Sono ruim (insônia / acorda cansada)", "group": "corpo"},
    {"id": "sym_fadiga", "label": "Cansaço / fadiga", "group": "corpo"},
    {"id": "sym_sobrecarga", "label": "Sobrecarga / excesso de responsabilidades", "group": "corpo"},

    # Mente
    {"id": "sym_ansiedade", "label": "Ansiedade / agitação", "group": "mente"},
    {"id": "sym_ruminacao", "label": "Pensamentos repetitivos / ruminação", "group": "mente"},
    {"id": "sym_irritabilidade", "label": "Irritabilidade / impaciência", "group": "mente"},

    # Pertencimento
    {"id": "sym_falta_pertenc", "label": "Falta de pertencimento / desconexão", "group": "pertenc"},
    {"id": "sym_autocritica", "label": "Autojulgamento / vergonha", "group": "pertenc"},
    {"id": "sym_isolamento", "label": "Isolamento / dificuldade de pedir ajuda", "group": "pertenc"},
]

V3_NEURO = [
    # 0–4 (quanto maior, melhor) exceto 'alerta' (quanto maior, pior)
    {"id": "np_alerta", "label": "Meu corpo/mente fica em alerta com facilidade", "invert": True},
    {"id": "np_pausa", "label": "Consigo pausar 30–60s antes de reagir", "invert": False},
    {"id": "np_recurso", "label": "Tenho pelo menos 1 recurso que me acalma (respiração, oração, música etc.)", "invert": False},
    {"id": "np_adesao", "label": "Consigo praticar um exercício curto todos os dias", "invert": False},
]

V3_TIME_OPTIONS = [
    (2, "2 min/dia"),
    (5, "5 min/dia"),
    (10, "10 min/dia"),
    (15, "15 min/dia"),
]

def _clamp_int(x: Any, lo: int, hi: int) -> int:
    try:
        xi = int(float(x))
    except Exception:
        xi = lo
    return lo if xi < lo else (hi if xi > hi else xi)

def _pct_from_0_10(v: int) -> float:
    v = _clamp_int(v, 0, 10)
    return round((v / 10.0) * 100.0, 1)

def _pct_from_0_4(v: int) -> float:
    v = _clamp_int(v, 0, 4)
    return round((v / 4.0) * 100.0, 1)

def compute_core_pillars(sym: Dict[str, int]) -> Dict[str, float]:
    """4 pilares principais (0–100)."""
    dor = _pct_from_0_10(sym.get("sym_dor", 0))
    rig = _pct_from_0_10(sym.get("sym_rigidez", 0))
    sono = _pct_from_0_10(sym.get("sym_sono", 0))
    fad  = _pct_from_0_10(sym.get("sym_fadiga", 0))
    sob  = _pct_from_0_10(sym.get("sym_sobrecarga", 0))

    ans  = _pct_from_0_10(sym.get("sym_ansiedade", 0))
    rum  = _pct_from_0_10(sym.get("sym_ruminacao", 0))
    irr  = _pct_from_0_10(sym.get("sym_irritabilidade", 0))

    per  = _pct_from_0_10(sym.get("sym_falta_pertenc", 0))
    aj   = _pct_from_0_10(sym.get("sym_autocritica", 0))
    iso  = _pct_from_0_10(sym.get("sym_isolamento", 0))

    # Pilares (médias ponderadas leves)
    return {
        "Dor/Fibro": round((dor * 0.7 + rig * 0.3), 1),
        "Sistema nervoso": round((ans * 0.45 + rum * 0.35 + irr * 0.20), 1),
        "Burnout/Cansaço": round((fad * 0.45 + sob * 0.35 + sono * 0.20), 1),
        "Pertencimento": round((per * 0.45 + aj * 0.35 + iso * 0.20), 1),
    }

def compute_readiness_pct(neuro: Dict[str, int], time_min: int) -> float:
    """Prontidão para neuroplasticidade (0–100)."""
    # itens positivos
    pausa = _clamp_int(neuro.get("np_pausa", 0), 0, 4)
    recurso = _clamp_int(neuro.get("np_recurso", 0), 0, 4)
    adesao = _clamp_int(neuro.get("np_adesao", 0), 0, 4)
    base = (pausa + recurso + adesao) / 12.0  # 0..1

    # alerta (negativo)
    alerta = _clamp_int(neuro.get("np_alerta", 0), 0, 4)  # alto = pior
    mod_alerta = 1.0 - (alerta / 6.0)  # 0.33..1
    if mod_alerta < 0.33:
        mod_alerta = 0.33

    # tempo disponível (pequeno bônus)
    time_min = _clamp_int(time_min, 2, 15)
    time_bonus = 1.0
    if time_min >= 10:
        time_bonus = 1.08
    elif time_min >= 5:
        time_bonus = 1.04

    pct = base * mod_alerta * time_bonus * 100.0
    return round(0.0 if pct < 0 else (100.0 if pct > 100 else pct), 1)

def compute_scores_v3(sym: Dict[str, int]) -> Tuple[Dict[str, float], Dict[str, float]]:
    """Converte o mapa 0–10 em scores por domínio existente (DOMAINS), e também devolve os pilares (4)."""
    pillars = compute_core_pillars(sym)

    # Domínios existentes (0–100)
    sono = _pct_from_0_10(sym.get("sym_sono", 0))
    ansiedade = round(pillars["Sistema nervoso"], 1)
    ruminacao = _pct_from_0_10(sym.get("sym_ruminacao", 0))
    exaustao = round(pillars["Burnout/Cansaço"], 1)
    pertenc = round(pillars["Pertencimento"], 1)
    tensao = round(pillars["Dor/Fibro"], 1)

    # humor baixo: aproximação (não diagnostica) — combina exaustão + pertencimento + sono
    humor = round((exaustao * 0.45 + pertenc * 0.35 + sono * 0.20), 1)

    scores = {
        "sono": float(sono),
        "ansiedade": float(ansiedade),
        "humor_baixo": float(humor),
        "exaustao": float(exaustao),
        "pertencimento": float(pertenc),
        "tensao": float(tensao),
        "ruminacao": float(ruminacao),
    }
    # garante todas as chaves de DOMAINS
    for d in DOMAINS:
        scores.setdefault(d, 0.0)
    return scores, pillars

def apply_neuro_to_sessions(qty: int, cadence: int, readiness_pct: float, dor_0_10: int) -> Tuple[int, int, List[str]]:
    """Ajusta quantidade/cadência de sessões conforme prontidão e dor (sem travar o terapeuta)."""
    qty_i = int(qty or 0)
    cad_i = int(cadence or 0)
    notes: List[str] = []

    dor_0_10 = _clamp_int(dor_0_10, 0, 10)

    # dor alta: mais suporte e cadência menor
    if dor_0_10 >= 7:
        qty_i += 1
        cad_i = max(7, cad_i - 2)
        notes.append("Dor alta → +1 sessão e cadência um pouco mais curta.")
    elif dor_0_10 >= 4:
        cad_i = max(7, cad_i - 1)
        notes.append("Dor moderada → cadência ligeiramente mais curta.")

    # prontidão baixa: manter mais próximo e mais suporte
    if readiness_pct < 35:
        qty_i += 2
        cad_i = max(7, cad_i - 3)
        notes.append("Prontidão baixa (neuroplasticidade) → +2 sessões e cadência mais curta.")
    elif readiness_pct < 55:
        qty_i += 1
        cad_i = max(7, cad_i - 1)
        notes.append("Prontidão média-baixa → +1 sessão e cadência levemente menor.")
    elif readiness_pct > 80:
        # prontidão alta: pode espaçar um pouco (se não houver dor alta)
        if dor_0_10 <= 3:
            cad_i = min(21, cad_i + 2)
            notes.append("Prontidão alta → pode espaçar um pouco a cadência.")
    qty_i = 1 if qty_i < 1 else (12 if qty_i > 12 else qty_i)
    cad_i = 7 if cad_i < 7 else (30 if cad_i > 30 else cad_i)
    return qty_i, cad_i, notes

QUESTIONS = [
    {"id": "sono_q1", "label": "Dificuldade para pegar no sono", "domain": "sono", "weight": 1.0},
    {"id": "sono_q2", "label": "Acorda no meio da noite / sono leve", "domain": "sono", "weight": 1.0},
    {"id": "ans_q1", "label": "Ansiedade / agitação no dia a dia", "domain": "ansiedade", "weight": 1.2},
    {"id": "ans_q2", "label": "Sintomas físicos de ansiedade (aperto, inquietação)", "domain": "ansiedade", "weight": 1.0},
    {"id": "hum_q1", "label": "Tristeza / desânimo frequente", "domain": "humor_baixo", "weight": 1.2},
    {"id": "hum_q2", "label": "Perda de prazer / motivação", "domain": "humor_baixo", "weight": 1.0},
    {"id": "exa_q1", "label": "Cansaço / exaustão por responsabilidades", "domain": "exaustao", "weight": 1.2},
    {"id": "exa_q2", "label": "Pouco tempo para si / autocuidado", "domain": "exaustao", "weight": 1.0},
    {"id": "per_q1", "label": "Sensação de não pertencimento / desconexão", "domain": "pertencimento", "weight": 1.2},
    {"id": "per_q2", "label": "Vergonha / autojulgamento", "domain": "pertencimento", "weight": 1.0},
    {"id": "ten_q1", "label": "Tensão muscular / dores recorrentes", "domain": "tensao", "weight": 1.0},
    {"id": "ten_q2", "label": "Mandíbula/ombros travados / corpo em alerta", "domain": "tensao", "weight": 1.0},
    {"id": "rum_q1", "label": "Mente acelerada / ruminação", "domain": "ruminacao", "weight": 1.2},
    {"id": "rum_q2", "label": "Dificuldade de foco por pensamentos repetitivos", "domain": "ruminacao", "weight": 1.0},

    {"id": "ans_q3", "label": "Ansiedade no momento (sensação presente)", "domain": "ansiedade", "weight": 0.8},
    {"id": "hum_q3", "label": "Histórico ou sinais atuais de depressão", "domain": "humor_baixo", "weight": 1.3},
    {"id": "exa_q3", "label": "Baixa energia hoje / esgotamento físico", "domain": "exaustao", "weight": 1.1},
    {"id": "ten_q3", "label": "Dor/incômodo físico atualmente", "domain": "tensao", "weight": 1.2},
    {"id": "ten_q4", "label": "Dor recorrente ou crônica (últimos meses)", "domain": "tensao", "weight": 1.0},
    {"id": "ten_q5", "label": "A dor limita atividades ou movimentos", "domain": "tensao", "weight": 1.0},
]

FLAGS = [
    {"id": "flag_preg", "label": "Gestação / amamentação"},
    {"id": "flag_meds", "label": "Uso de medicamentos (ansiolíticos/antidepressivos/sedativos)"},
    {"id": "flag_allergy", "label": "Alergias / sensibilidades"},
    {"id": "flag_sound", "label": "Sensibilidade a som (binaural)"},
    {"id": "flag_light", "label": "Sensibilidade à luz (cama de cristal)"},

    {"id": "flag_back", "label": "Dificuldade para deitar de costas (cama de cristal)"},
    {"id": "flag_heat", "label": "Sente muito calor / sensibilidade ao calor"},
    {"id": "flag_perfume", "label": "Sensibilidade a cheiros/perfumes (aromaterapia)"},
    {"id": "flag_feet", "label": "Sensibilidade nos pés (pressão/reflexo)"},
]

DOMAIN_TO_PROTOCOL = {
    "ansiedade": "FOCO – Ansiedade / Agitação",
    "sono": "FOCO – Sono Profundo",
    "exaustao": "FOCO – Exaustão / Sobrecarga",
    "pertencimento": "FOCO – Pertencimento / Vergonha",
    "humor_baixo": "FOCO – Exaustão / Sobrecarga",
    "tensao": "FOCO – Ansiedade / Agitação",
    "ruminacao": "FOCO – Ansiedade / Agitação",
}
BASE_PROTOCOL = "BASE – Aterramento + Regulação"


def compute_scores(answers: Dict[str, int]) -> Dict[str, float]:
    sums = {d: 0.0 for d in DOMAINS}
    maxs = {d: 0.0 for d in DOMAINS}
    for q in QUESTIONS:
        v = float(answers.get(q["id"], 0))
        w = float(q["weight"])
        d = q["domain"]
        sums[d] += v * w
        maxs[d] += 4.0 * w
    return {d: round((sums[d] / maxs[d] * 100.0) if maxs[d] else 0.0, 1) for d in DOMAINS}



def adjust_scores_with_phys(scores: Dict[str, int], phys_meta: Dict[str, Any]) -> Tuple[Dict[str, int], Dict[str, Any]]:
    """Ajustes leves de score com base em dor/contexto (sem 'diagnosticar').

    Retorna: (scores_ajustados, contexto) onde contexto inclui alertas/tags.
    """
    out = {k: int(v or 0) for k, v in (scores or {}).items()}
    ctx_alertas: List[str] = []
    ctx_tags: List[str] = []
    ctx_ajustes: List[str] = []

    def clamp100(x: int) -> int:
        return 0 if x < 0 else (100 if x > 100 else int(x))

    def add(domain: str, delta: int, reason: str = ""):
        if domain not in out:
            return
        before = out[domain]
        out[domain] = clamp100(before + int(delta))
        if reason and out[domain] != before:
            ctx_ajustes.append(f"{domain}: {before}→{out[domain]} ({reason})")

    pm = phys_meta or {}

    # Dor (0–10)
    dor = pm.get("phys_dor_score")
    try:
        dor_i = int(float(dor)) if dor is not None else 0
    except Exception:
        dor_i = 0

    if dor_i >= 7:
        add("tensao", 12, "dor intensa")
        add("exaustao", 8, "dor intensa")
        ctx_tags.append("dor_intensa")
    elif dor_i >= 4:
        add("tensao", 6, "dor moderada")
        add("exaustao", 4, "dor moderada")
        ctx_tags.append("dor_moderada")

    # Emoções (auto-relato)
    emo = (pm.get("phys_emocoes_lida") or "").strip()
    if emo == "Guardo pra mim / engulo":
        add("ruminacao", 6, "tende a engolir emoções")
        add("ansiedade", 4, "tende a engolir emoções")
        ctx_tags.append("emocao_internaliza")
    elif emo == "Explodo / fico irritada":
        add("ansiedade", 6, "tende a explodir")
        add("tensao", 4, "tende a explodir")
        ctx_tags.append("emocao_reatividade")
    elif emo == "Choro / fico retraída":
        add("humor_baixo", 4, "tende a retração")
        ctx_tags.append("emocao_retracao")

    # Conflito familiar
    conf = (pm.get("phys_conflito_nivel") or "Não").strip()
    if conf == "Leve":
        add("pertencimento", 3, "conflito familiar leve")
        add("ansiedade", 3, "conflito familiar leve")
        ctx_tags.append("conflito_familiar")
    elif conf == "Moderado":
        add("pertencimento", 6, "conflito familiar moderado")
        add("ansiedade", 6, "conflito familiar moderado")
        add("ruminacao", 4, "conflito familiar moderado")
        ctx_tags.append("conflito_familiar")
    elif conf == "Grave":
        add("pertencimento", 10, "conflito familiar grave")
        add("ansiedade", 10, "conflito familiar grave")
        add("ruminacao", 6, "conflito familiar grave")
        ctx_tags.append("conflito_familiar")

    # Transtorno alimentar (auto-relato)
    ta = (pm.get("phys_transt_alim") or "Não").strip()
    if ta == "Suspeita/Em investigação":
        add("ansiedade", 6, "transtorno alimentar (suspeita)")
        add("humor_baixo", 6, "transtorno alimentar (suspeita)")
        ctx_tags.append("transtorno_alimentar")
    elif ta == "Sim":
        add("ansiedade", 10, "transtorno alimentar (sim)")
        add("humor_baixo", 10, "transtorno alimentar (sim)")
        ctx_tags.append("transtorno_alimentar")

    # Alertas (sem mexer em score)
    if (pm.get("phys_alergias") or "Não") == "Sim":
        quais = (pm.get("phys_alergias_quais") or "").strip()
        ctx_alertas.append("⚠️ Alergias relatadas" + (f": {quais}" if quais else "."))
        ctx_tags.append("alergias")

    if (pm.get("phys_cirurgias") or "Não") == "Sim":
        quais = (pm.get("phys_cirurgias_quais") or "").strip()
        ctx_alertas.append("⚠️ Cirurgias relatadas" + (f": {quais}" if quais else "."))
        ctx_tags.append("cirurgias")

    fam = (pm.get("phys_hist_familia") or "").strip()
    if fam:
        ctx_alertas.append("ℹ️ Histórico familiar relevante informado.")
        ctx_tags.append("hist_familiar")

    # Dedup
    ctx_alertas = list(dict.fromkeys([a for a in ctx_alertas if a]))
    ctx_tags = list(dict.fromkeys([t for t in ctx_tags if t]))

    contexto = {
        "alertas": ctx_alertas,
        "tags": ctx_tags,
        "ajustes": ctx_ajustes,
    }
    return out, contexto

def pick_focus(scores: Dict[str, float], top_n: int = 3) -> List[Tuple[str, float]]:
    return sorted(scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

# -----------------------------
# V4 — Atendimento PRO: 9 Origens (triagem integrativa)
# Objetivo: gerar um raciocínio mais "clínico" (com evidências + medição + contexto),
# e ainda assim manter compatibilidade com DOMAINS/protocolos existentes.
# -----------------------------
ORIGENS9 = [
    {
        "id": "psicoemocional",
        "label": "Psicoemocional (estresse / emoções)",
        "desc": "Quando emoções não processadas e estresse sustentado viram sintoma no corpo/mente.",
        "evidence": [
            {"id": "piora_estresse", "label": "Piora clara em períodos de estresse/pressão", "w": 2.0},
            {"id": "irrit_choro", "label": "Irritabilidade, choro fácil ou sensibilidade aumentada", "w": 1.5},
            {"id": "carga_cuidador", "label": "Excesso de responsabilidades / cuida de alguém", "w": 1.5},
            {"id": "luto_trauma", "label": "Histórico de luto/trauma/choques emocionais", "w": 1.2},
            {"id": "somatiza", "label": "Sente que 'segura' tudo e somatiza no corpo", "w": 1.0},
        ],
    },
    {
        "id": "mental_cognitivo",
        "label": "Mental / Cognitivo (crenças / ruminação)",
        "desc": "Padrões mentais repetitivos, autocobrança e hipercontrole mantendo o sistema em tensão.",
        "evidence": [
            {"id": "ruminacao", "label": "Pensamentos repetitivos (difícil 'desligar')", "w": 2.0},
            {"id": "autocritica", "label": "Autojulgamento / perfeccionismo / vergonha", "w": 1.5},
            {"id": "catastrofe", "label": "Tende a antecipar problemas / catastrofizar", "w": 1.0},
            {"id": "foco_baixo", "label": "Dificuldade de foco e memória por mente cheia", "w": 1.0},
        ],
    },
    {
        "id": "energetico_espiritual",
        "label": "Energético / Espiritual (campo / pertencimento)",
        "desc": "Desalinhamentos sutis, falta de aterramento e 'desconexão' da própria essência.",
        "evidence": [
            {"id": "sem_pertenc", "label": "Falta de pertencimento / desconexão de si", "w": 2.0},
            {"id": "cansaço_amb", "label": "Cansaço após ambientes/pessoas (sensibilidade)", "w": 1.5},
            {"id": "sono_sensivel", "label": "Sono sensível a energia/ambiente (acorda 'carregada')", "w": 1.0},
            {"id": "precisa_limpeza", "label": "Sente necessidade frequente de limpeza/ritual", "w": 0.8},
        ],
    },
    {
        "id": "neurovegetativo",
        "label": "Sistema nervoso / Neurovegetativo (alerta)",
        "desc": "Hiperalerta, disautonomia leve e baixa capacidade de recuperação (stress response).",
        "evidence": [
            {"id": "ansiedade", "label": "Ansiedade/agitação, 'corpo em alerta'", "w": 2.0},
            {"id": "sono_ruim", "label": "Sono não reparador / despertares", "w": 1.5},
            {"id": "sens_ruido", "label": "Sensível a barulho/luz/toques, irrita fácil", "w": 1.0},
            {"id": "tensao", "label": "Tensão muscular / mandíbula / ombros", "w": 1.0},
        ],
    },
    {
        "id": "inflamatorio_imune",
        "label": "Inflamatório / Imunológico (dor / inflamação)",
        "desc": "Processos inflamatórios (agudos ou crônicos) sustentando dor, rigidez e fadiga.",
        "evidence": [
            {"id": "dor_rigidez", "label": "Dor + rigidez/travamento frequentes", "w": 2.0},
            {"id": "piora_clima", "label": "Piora com frio/umidade ou após esforço", "w": 1.0},
            {"id": "pele_sinus", "label": "Quadros alérgicos (rinite/sinus/pele)", "w": 1.0},
            {"id": "inflam_dieta", "label": "Piora com certos alimentos (açúcar, ultraprocessados)", "w": 1.0},
        ],
    },
    {
        "id": "bioquimico_nutricional",
        "label": "Bioquímico / Nutricional (energia / carências)",
        "desc": "Rotina alimentar, hidratação e possíveis carências afetando energia e humor.",
        "evidence": [
            {"id": "picos_acucar", "label": "Oscilações de energia com açúcar/café (picos e quedas)", "w": 1.5},
            {"id": "pouca_agua", "label": "Hidratação baixa", "w": 1.0},
            {"id": "pouca_proteina", "label": "Baixa proteína/fibras na rotina", "w": 1.0},
            {"id": "rotina_irreg", "label": "Rotina irregular (pula refeições / come tarde)", "w": 1.2},
        ],
    },
    {
        "id": "toxinas",
        "label": "Tóxico / Químico (sobrecarga)",
        "desc": "Exposição a químicos, fumaça, excesso de remédios/suplementos ou ambiente poluído.",
        "evidence": [
            {"id": "cheiros", "label": "Sensível a cheiros (perfumes/limpeza)", "w": 1.2},
            {"id": "amb_poluido", "label": "Ambiente com fumaça/poluição/mofo", "w": 1.2},
            {"id": "muitos_meds", "label": "Uso de muitos medicamentos/suplementos", "w": 1.0},
            {"id": "cabeca", "label": "Dor de cabeça/névoa mental recorrente", "w": 1.0},
        ],
    },
    {
        "id": "microbiota_infeccioso",
        "label": "Microbiota / Infeccioso (intestino)",
        "desc": "Disbiose, intestino irregular e cargas infecciosas afetando corpo e mente.",
        "evidence": [
            {"id": "intestino", "label": "Intestino irregular (prisão/diarreia/alternância)", "w": 2.0},
            {"id": "inchaço", "label": "Inchaço/gases após comer", "w": 1.2},
            {"id": "antibiotico", "label": "Uso recente de antibiótico ou muitas infecções", "w": 1.0},
            {"id": "refluxo", "label": "Refluxo/azia recorrentes", "w": 1.0},
        ],
    },
    {
        "id": "estrutural",
        "label": "Estrutural / Funcional (corpo físico)",
        "desc": "Postura, lesões, sedentarismo, falta de mobilidade e sobrecarga mecânica.",
        "evidence": [
            {"id": "dor_local", "label": "Dor bem localizada (coluna, cervical, ombro etc.)", "w": 2.0},
            {"id": "sedentarismo", "label": "Pouco movimento / alongamento", "w": 1.2},
            {"id": "postura", "label": "Trabalha muito sentada(o) / postura ruim", "w": 1.0},
            {"id": "bruxismo", "label": "Bruxismo / mandíbula travada", "w": 0.8},
        ],
    },
]

# Como as origens alimentam DOMAINS (mantém compatibilidade com a biblioteca de protocolos)
ORIGEM_TO_DOMAIN_WEIGHTS: Dict[str, Dict[str, float]] = {
    "psicoemocional": {"ansiedade": 0.8, "humor_baixo": 0.6, "pertencimento": 0.7, "ruminacao": 0.5, "sono": 0.5, "tensao": 0.3, "exaustao": 0.4},
    "mental_cognitivo": {"ruminacao": 0.9, "humor_baixo": 0.5, "ansiedade": 0.4, "pertencimento": 0.4, "sono": 0.3},
    "energetico_espiritual": {"pertencimento": 0.9, "ansiedade": 0.4, "humor_baixo": 0.3, "sono": 0.2},
    "neurovegetativo": {"ansiedade": 0.7, "sono": 0.8, "tensao": 0.5, "exaustao": 0.6, "ruminacao": 0.5},
    "inflamatorio_imune": {"tensao": 0.7, "exaustao": 0.5, "sono": 0.3, "humor_baixo": 0.3},
    "bioquimico_nutricional": {"exaustao": 0.7, "sono": 0.3, "humor_baixo": 0.3, "ansiedade": 0.2},
    "toxinas": {"exaustao": 0.6, "humor_baixo": 0.4, "ansiedade": 0.3, "sono": 0.2},
    "microbiota_infeccioso": {"exaustao": 0.6, "humor_baixo": 0.3, "ansiedade": 0.2, "sono": 0.2},
    "estrutural": {"tensao": 0.8, "sono": 0.2, "exaustao": 0.2},
}

# Sintomas (0–10) também alimentam DOMAINS — dá variação real e evita "sempre a mesma resposta"
SYM_TO_DOMAIN: Dict[str, Dict[str, float]] = {
    "sym_sono": {"sono": 1.0},
    "sym_ansiedade": {"ansiedade": 1.0, "ruminacao": 0.2},
    "sym_ruminacao": {"ruminacao": 1.0},
    "sym_fadiga": {"exaustao": 0.8, "sono": 0.2},
    "sym_sobrecarga": {"exaustao": 0.6, "ansiedade": 0.2},
    "sym_dor": {"tensao": 0.7, "sono": 0.1},
    "sym_rigidez": {"tensao": 0.7},
    "sym_falta_pertenc": {"pertencimento": 1.0},
    "sym_isolamento": {"pertencimento": 0.6, "humor_baixo": 0.3},
    "sym_autocritica": {"humor_baixo": 0.6, "pertencimento": 0.4, "ruminacao": 0.2},
    "sym_irritabilidade": {"ansiedade": 0.5, "humor_baixo": 0.2, "tensao": 0.2},
}

def _softmax(xs: List[float], temp: float = 1.0) -> List[float]:
    if not xs:
        return []
    t = float(temp or 1.0)
    t = 1.0 if t <= 0 else t
    m = max(xs)
    exps = [math.exp((x - m) / t) for x in xs]
    s = sum(exps) or 1.0
    return [e / s for e in exps]

def _conf_label(c: float) -> str:
    if c >= 0.72:
        return "Alta"
    if c >= 0.52:
        return "Média"
    return "Baixa"

def compute_origens9(
    measures_0_10: Dict[str, float],
    evidences: Dict[str, Dict[str, bool]],
    sym_0_10: Dict[str, int],
) -> Dict[str, Any]:
    """Retorna pontuação (0–100), probabilidade (0–1) e confiança (0–1) por origem.
    - measures_0_10: medição manual (aurímetro/pêndulo/escuta) por origem (0–10)
    - evidences: {orig_id: {evidence_id: bool}}
    - sym_0_10: mapa rápido 0–10
    """
    raw: Dict[str, float] = {}
    conf: Dict[str, float] = {}
    hits: Dict[str, List[str]] = {}

    # contexto via sintomas (até 10 pontos)
    def ctx_for_origin(orig_id: str) -> float:
        # soma contribuições de sintomas relevantes para a origem (normalizado)
        # mapeamento simples: usa SYM_TO_DOMAIN como proxy + ORIGEM_TO_DOMAIN_WEIGHTS
        w_dom = ORIGEM_TO_DOMAIN_WEIGHTS.get(orig_id, {})
        ctx = 0.0
        for sid, val in (sym_0_10 or {}).items():
            v = float(val or 0) / 10.0
            if v <= 0:
                continue
            # converte sintoma->domínio->origem (aproximação consistente)
            for dom, wsd in (SYM_TO_DOMAIN.get(sid, {}) or {}).items():
                ctx += v * float(wsd) * float(w_dom.get(dom, 0.0))
        # comprime: 0..~? para 0..10
        return max(0.0, min(10.0, ctx * 6.0))

    for o in ORIGENS9:
        oid = o["id"]
        m = float(measures_0_10.get(oid, 0.0) or 0.0)
        m = max(0.0, min(10.0, m))
        m_norm = m / 10.0

        # evidências
        ev_list = o.get("evidence") or []
        ev_map = evidences.get(oid, {}) if isinstance(evidences.get(oid, {}), dict) else {}
        total_w = sum(float(e.get("w", 1.0)) for e in ev_list) or 1.0
        got_w = 0.0
        hit_labels: List[str] = []
        for e in ev_list:
            eid = e.get("id")
            if not eid:
                continue
            if bool(ev_map.get(eid, False)):
                got_w += float(e.get("w", 1.0))
                hit_labels.append(str(e.get("label") or eid))
        ev_norm = got_w / total_w  # 0..1

        ctx = ctx_for_origin(oid)  # 0..10
        ctx_norm = ctx / 10.0

        # score: medição (0..60) + evidência (0..30) + contexto (0..10)
        score = (m_norm * 60.0) + (ev_norm * 30.0) + (ctx_norm * 10.0)
        score = max(0.0, min(100.0, score))
        raw[oid] = score
        hits[oid] = hit_labels

        # confiança: evidência + consistência medição/evidência + presença de contexto
        consistency = 1.0 - abs(m_norm - ev_norm)  # 0..1
        c = 0.30 + 0.38 * ev_norm + 0.22 * consistency + 0.10 * ctx_norm
        conf[oid] = max(0.0, min(1.0, c))

    # probabilidade normalizada (softmax para evitar empate constante)
    probs = _softmax([raw[o["id"]] for o in ORIGENS9], temp=12.0)
    prob: Dict[str, float] = {o["id"]: float(p) for o, p in zip(ORIGENS9, probs)}

    return {"raw": raw, "prob": prob, "confidence": conf, "hits": hits}

def compute_domain_scores_prof(
    orig_prob: Dict[str, float],
    sym_0_10: Dict[str, int],
    phys_meta: Dict[str, Any],
) -> Dict[str, float]:
    """Combina origens (prob) + sintomas (0–10) + físico (pouco) => scores 0–100 por DOMAINS."""
    # 1) contribuição das origens (até 60)
    dom_from_orig = {d: 0.0 for d in DOMAINS}
    for oid, p in (orig_prob or {}).items():
        wmap = ORIGEM_TO_DOMAIN_WEIGHTS.get(oid, {})
        for dom, w in wmap.items():
            if dom in dom_from_orig:
                dom_from_orig[dom] += float(p) * float(w)
    # normaliza por "máximo teórico" aproximado
    for dom in dom_from_orig:
        dom_from_orig[dom] = max(0.0, min(1.0, dom_from_orig[dom])) * 60.0

    # 2) contribuição dos sintomas (até 40)
    dom_from_sym = {d: 0.0 for d in DOMAINS}
    for sid, val in (sym_0_10 or {}).items():
        v = max(0.0, min(10.0, float(val or 0.0))) / 10.0
        for dom, w in (SYM_TO_DOMAIN.get(sid, {}) or {}).items():
            if dom in dom_from_sym:
                dom_from_sym[dom] += v * float(w)
    for dom in dom_from_sym:
        dom_from_sym[dom] = max(0.0, min(1.0, dom_from_sym[dom])) * 40.0

    scores = {dom: max(0.0, min(100.0, dom_from_orig.get(dom, 0.0) + dom_from_sym.get(dom, 0.0))) for dom in DOMAINS}

    # 3) ajuste físico leve (evita respostas iguais sem virar diagnóstico)
    try:
        dor = float(phys_meta.get("phys_dor_score") or 0.0) / 10.0
        if dor > 0:
            scores["tensao"] = max(scores["tensao"], min(100.0, scores["tensao"] + dor * 12.0))
    except Exception:
        pass
    return scores


def sessions_from_scores(scores: Dict[str, float]) -> Tuple[int, int]:
    top = sorted(scores.values(), reverse=True)
    max_score = top[0] if top else 0.0
    strong = sum(1 for s in top[:3] if s >= 70)
    if strong <= 0:
        strong = sum(1 for s in top[:3] if s >= 60)
    qty = 4 if strong <= 1 else (6 if strong == 2 else 8)
    cadence = 7 if max_score >= 80 else (10 if max_score >= 60 else 14)
    return qty, cadence




def build_alert_actions(
    alertas: List[str],
    flags: Dict[str, Any],
    phys_meta: Dict[str, Any],
    ctx_phys: Dict[str, Any],
) -> List[Dict[str, str]]:
    """Transforma alertas em condutas práticas (para entrar no plano)."""
    out: List[Dict[str, str]] = []

    def add(categoria: str, detalhe: str, acao: str, prioridade: str = "Média"):
        out.append({
            "Prioridade": prioridade,
            "Categoria": categoria,
            "Detalhe": (detalhe or "").strip(),
            "Conduta sugerida": acao.strip(),
        })

    # Depressão (0-4) — do físico
    dep_lvl = int(phys_meta.get("phys_depressao_nivel") or 0)
    if dep_lvl >= 3:
        add(
            "Humor / Depressão",
            f"Nível {dep_lvl}/4",
            "Priorizar acolhimento e regulação suave (α/θ). Investigar sinais de risco e sugerir acompanhamento com psicólogo/psiquiatra quando necessário. Evitar estímulos intensos nas primeiras sessões.",
            "Alta",
        )
    elif dep_lvl == 2:
        add(
            "Humor / Depressão",
            "Nível 2/4 (moderada)",
            "Aumentar foco em regulação do SNC, rotina mínima e suporte emocional. Sugerir acompanhamento profissional se houver piora ou ideação.",
            "Média-Alta",
        )

    # Conflito familiar
    confl = str(phys_meta.get("phys_conflito_nivel") or "").strip()
    if confl in ("Moderado", "Grave"):
        add(
            "Conflito familiar",
            confl,
            "Incluir práticas de segurança interna (grounding/acolhimento), pertencimento e limites. Sugerir rede de apoio/terapia. Se houver violência/ameaça, orientar buscar ajuda imediata (serviços locais).",
            "Alta" if confl == "Grave" else "Média",
        )

    # Transtorno alimentar
    ta = str(phys_meta.get("phys_transt_alim") or "").strip()
    if ta in ("Sim", "Suspeita/Em investigação"):
        add(
            "Transtorno alimentar",
            ta,
            "Evitar intervenções centradas em peso/culpa. Priorizar regulação emocional e autoimagem. Sugerir acompanhamento conjunto com nutricionista e psicólogo/psiquiatra.",
            "Alta" if ta == "Sim" else "Média",
        )

    # Alergias / sensibilidades
    if str(phys_meta.get("phys_alergias") or "") == "Sim" or bool(flags.get("flag_allergy")):
        quais = (phys_meta.get("phys_alergias_quais") or "").strip()
        add(
            "Alergias / sensibilidades",
            quais or "Relatada na anamnese",
            "Registrar no prontuário. Evitar fitoterápicos/aromas potencialmente irritantes e iniciar com mínima exposição (teste/monitoramento). Em reações, suspender e orientar avaliação médica.",
            "Média",
        )

    # Cirurgias
    if str(phys_meta.get("phys_cirurgias") or "") == "Sim":
        quais = (phys_meta.get("phys_cirurgias_quais") or "").strip()
        add(
            "Cirurgias / condições prévias",
            quais or "Relatada na anamnese",
            "Respeitar limitações de postura/tempo na cama. Se cirurgia recente, dor intensa ou sinais de alerta, orientar avaliação médica antes de intensificar protocolos.",
            "Média",
        )

    # Dor (0-10)
    dor = int(phys_meta.get("phys_dor_score") or 0)
    if dor >= 8:
        add(
            "Dor",
            f"{dor}/10",
            "Dor intensa: avaliar sinais de alarme e sugerir avaliação médica/fisioterapia. Nas sessões, priorizar relaxamento/analgesia suave (α/θ) e luz baixa; progressão gradual.",
            "Alta",
        )
    elif dor >= 5:
        add(
            "Dor",
            f"{dor}/10",
            "Incluir abordagem para tensão/dor (relaxamento, respiração, alongamento leve). Monitorar melhora entre sessões e ajustar intensidade.",
            "Média",
        )

    # Medicamentos
    meds_txt = (phys_meta.get("phys_meds_txt") or "").strip()
    if meds_txt or bool(flags.get("flag_meds")):
        add(
            "Medicamentos",
            meds_txt or "Uso relatado",
            "Registrar medicamentos em uso. Evitar promessas terapêuticas e orientar que qualquer ajuste de medicação deve ser feito apenas com o médico.",
            "Baixa",
        )

    # Sensibilidade a som/luz
    if bool(flags.get("flag_sound")):
        add(
            "Sensibilidade a som",
            "",
            "Manter volume baixo, músicas neutras, evitar frequências mais estimulantes no início e checar conforto durante a sessão.",
            "Média",
        )
    if bool(flags.get("flag_light")):
        add(
            "Sensibilidade a luz",
            "",
            "Usar intensidade de luz baixa/moderada, evitar flashes e cores muito saturadas no início; checar conforto ocular/dor de cabeça.",
            "Média",
        )

    # Se existem alertas gerais do plano (flags), mas nada entrou acima, ainda assim registrar
    if alertas and not out:
        add("Alertas", "; ".join(alertas[:6]), "Registrar e considerar na condução das sessões (intensidade, acolhimento e encaminhamentos).", "Média")

    return out

def load_protocols() -> Dict[str, Dict[str, Any]]:
    if BACKEND == "postgres":
        rows = qall("select name, domain, rules_json, content_json, active from public.protocol_library where active = true")
        return {r["name"]: {"name": r["name"], "domain": r["domain"], "rules": r["rules_json"], "content": r["content_json"]} for r in rows}
    rows = sb_select("protocol_library", columns="name,domain,rules_json,content_json,active", eq={"active": True}, order=("domain", True))
    out: Dict[str, Dict[str, Any]] = {}
    for r in rows:
        out[r["name"]] = {"name": r["name"], "domain": r["domain"], "rules": r.get("rules_json") or {}, "content": r.get("content_json") or {}}
    return out


def load_binaural_presets() -> List[Dict[str, Any]]:
    if BACKEND == "postgres":
        return qall("select id, nome, carrier_hz, beat_hz, duracao_min, notas from public.binaural_presets order by nome")
    return sb_select("binaural_presets", columns="id,nome,carrier_hz,beat_hz,duracao_min,notas", order=("nome", True))


def load_frequencies(tipo: Optional[str] = None) -> List[Dict[str, Any]]:
    if BACKEND == "postgres":
        if tipo:
            return qall("select code, nome, hz, tipo, chakra, cor, descricao from public.frequencies where tipo=%s order by code", (tipo,))
        return qall("select code, nome, hz, tipo, chakra, cor, descricao from public.frequencies order by code")
    if tipo:
        return sb_select("frequencies", columns="code,nome,hz,tipo,chakra,cor,descricao", eq={"tipo": tipo}, order=("code", True))
    return sb_select("frequencies", columns="code,nome,hz,tipo,chakra,cor,descricao", order=("code", True))


def select_protocols(scores: Dict[str, float], protocols: Dict[str, Dict[str, Any]]) -> List[str]:
    selected = [BASE_PROTOCOL] if BASE_PROTOCOL in protocols else []
    for dom, sc in scores.items():
        if sc >= 60:
            pname = DOMAIN_TO_PROTOCOL.get(dom)
            if pname and pname in protocols and pname not in selected:
                selected.append(pname)
    if BASE_PROTOCOL in protocols and BASE_PROTOCOL not in selected:
        selected.insert(0, BASE_PROTOCOL)
    return selected


def merge_plan(selected_names: List[str], protocols: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    chakras: List[Any] = []
    emocoes: List[Any] = []
    cristais: List[Any] = []
    fito: List[Any] = []
    alerts: List[Any] = []

    def add_unique(lst, item):
        if item is None:
            return
        if item not in lst:
            lst.append(item)

    for name in selected_names:
        c = protocols.get(name, {}).get("content", {}) or {}
        for ch in c.get("chakras_foco", []):
            add_unique(chakras, ch)
        for e in c.get("emocoes_foco", []):
            add_unique(emocoes, e)
        for cr in c.get("cristais", []):
            add_unique(cristais, cr)
        for f in c.get("fito", []):
            add_unique(fito, f)
        if c.get("alertas"):
            add_unique(alerts, c["alertas"])
    return {
        "chakras_prioritarios": chakras,
        "emocoes_prioritarias": emocoes,
        "cristais_sugeridos": cristais,
        "fito_sugerida": fito,
        "alertas": alerts,
    }


def build_session_scripts(
    qty: int,
    cadence_days: int,
    focus: List[Tuple[str, float]],
    selected_names: List[str],
    protocols: Dict[str, Dict[str, Any]],
    audio_block: Dict[str, Any],
    extra_freq_codes: List[str],
) -> List[Dict[str, Any]]:
    focus_domains = [d for d, _ in focus]
    focus_cards: List[str] = []
    for dom in focus_domains:
        pname = DOMAIN_TO_PROTOCOL.get(dom)
        if pname and pname in selected_names and pname != BASE_PROTOCOL:
            focus_cards.append(pname)
    if not focus_cards:
        focus_cards = [n for n in selected_names if n != BASE_PROTOCOL][:1]

    # --- Sugestões consolidadas (cama de cristal + binaural dos protocolos) ---
    # (usadas para exibição na aba Atendimento e também salvas no script_json)
    cama_rows: List[Dict[str, Any]] = []
    proto_binaural_rows: List[Dict[str, Any]] = []

    def _add_protocol_suggestions(card_name: str, cama_obj: Any, binaural_obj: Any):
        # Cama de cristal
        if cama_obj is not None:
            if isinstance(cama_obj, list):
                for i, it in enumerate(cama_obj, start=1):
                    row = {"protocolo": card_name, "ordem": i}
                    if isinstance(it, dict):
                        row.update(it)
                    else:
                        row["item"] = str(it)
                    cama_rows.append(row)
            elif isinstance(cama_obj, dict):
                row = {"protocolo": card_name}
                row.update(cama_obj)
                cama_rows.append(row)
            else:
                cama_rows.append({"protocolo": card_name, "cama": str(cama_obj)})

        # Binaural do protocolo
        if binaural_obj is not None:
            if isinstance(binaural_obj, dict):
                row = {"protocolo": card_name}
                row.update(binaural_obj)
                proto_binaural_rows.append(row)
            else:
                proto_binaural_rows.append({"protocolo": card_name, "binaural": str(binaural_obj)})

    for _card in selected_names:
        c = protocols.get(_card, {}).get("content", {}) or {}
        _add_protocol_suggestions(_card, c.get("cama_cristal"), c.get("binaural"))

    scripts: List[Dict[str, Any]] = []
    today = date.today()

    for i in range(1, qty + 1):
        session_date = today + timedelta(days=cadence_days * (i - 1))
        focus_card = focus_cards[(i - 1) % len(focus_cards)] if focus_cards else None

        parts: List[Dict[str, Any]] = []
        base = protocols.get(BASE_PROTOCOL, {}).get("content", {}) or {}
        if base:
            parts.append(
                {
                    "card": BASE_PROTOCOL,
                    "binaural": base.get("binaural"),
                    "cama": base.get("cama_cristal"),
                    "cristais": base.get("cristais"),
                    "fito": base.get("fito"),
                    "roteiro": base.get("roteiro_sessao"),
                }
            )

        if focus_card:
            fc = protocols.get(focus_card, {}).get("content", {}) or {}
            parts.append(
                {
                    "card": focus_card,
                    "binaural": fc.get("binaural"),
                    "cama": fc.get("cama_cristal"),
                    "cristais": fc.get("cristais"),
                    "fito": fc.get("fito"),
                    "roteiro": fc.get("roteiro_sessao"),
                }
            )

        scripts.append(
            {
                "session_n": i,
                "scheduled_date": str(session_date),
                "status": "AGENDADA",
                "audio": audio_block,
                "frequencias": [{"code": c} for c in extra_freq_codes],
                "cama_cristal_sugestao": cama_rows,
                "binaural_protocolos_sugestao": proto_binaural_rows,
                "parts": parts,
            }
        )
    return scripts


# -------------------------
# Binaural: utilitários iguais ao app antigo
# -------------------------
MAX_BG_MB = 12  # ~12MB (vira ~16MB base64)

def bytes_to_data_url_safe(raw: Optional[bytes], filename: Optional[str], max_mb: int = MAX_BG_MB):
    """Converte bytes em data URL, mas recusa arquivos grandes para evitar MessageSizeError no Streamlit."""
    if not raw:
        return None, None, None
    size_mb = len(raw) / (1024 * 1024)
    name = (filename or "").lower()
    mime = "audio/mpeg"
    if name.endswith(".wav"):
        mime = "audio/wav"
    elif name.endswith(".ogg") or name.endswith(".oga"):
        mime = "audio/ogg"

    if size_mb > max_mb:
        return None, mime, f"Arquivo de {size_mb:.1f} MB excede o limite de {max_mb} MB para tocar embutido. Use arquivo menor."
    b64 = base64.b64encode(raw).decode("ascii")
    return f"data:{mime};base64,{b64}", mime, None


def synth_binaural_wav(carrier_hz: float, beat_hz: float, seconds: int = 20, sr: int = 44100, amp: float = 0.2) -> bytes:
    t = np.linspace(0, seconds, int(sr * seconds), endpoint=False)
    l = np.sin(2 * np.pi * (carrier_hz - beat_hz / 2.0) * t)
    r = np.sin(2 * np.pi * (carrier_hz + beat_hz / 2.0) * t)
    stereo = np.stack([l, r], axis=1) * float(amp)
    stereo_i16 = np.int16(np.clip(stereo, -1, 1) * 32767)

    bio = io.BytesIO()
    with wave.open(bio, "wb") as wf:
        wf.setnchannels(2)
        wf.setsampwidth(2)
        wf.setframerate(sr)
        wf.writeframes(stereo_i16.tobytes())
    return bio.getvalue()


def webaudio_binaural_html(
    fc: float,
    beat: float,
    seconds: int = 60,
    bg_data_url: Optional[str] = None,
    bg_gain: float = 0.12,
    binaural_gain: float = 0.20,
):
    """Player binaural + música de fundo usando <audio> + WebAudio (com botões Tocar/Parar)."""
    bt = abs(float(beat))
    fl = max(20.0, float(fc) - bt / 2)
    fr = float(fc) + bt / 2
    sec = int(max(5, seconds))
    bg = json.dumps(bg_data_url) if bg_data_url else "null"
    g = float(bg_gain)
    tg = float(binaural_gain)

    return f"""
<div style=\"padding:.6rem;border:1px solid #eee;border-radius:10px;\">
  <b>Binaural</b> — L {fl:.2f} Hz • R {fr:.2f} Hz • {sec}s {'<span style="margin-left:6px;">🎵 fundo</span>' if bg_data_url else ''}<br/>
  <button id=\"bplay\">▶️ Tocar</button> <button id=\"bstop\">⏹️ Parar</button>
  <div style=\"font-size:.9rem;color:#666\">Use fones · volume moderado</div>
</div>
<script>
let ctx=null, l=null, r=null, gL=null, gR=null, merger=null, bus=null, limiter=null, timer=null;
let bgAudio=null, bgNode=null, bgGain=null;

function cleanup(){{
  try{{ if(l) l.stop(); if(r) r.stop(); }}catch(e){{}}
  [l,r,gL,gR,merger,bus,limiter].forEach(n=>{{ if(n) try{{ n.disconnect(); }}catch(_e){{}} }});
  if(bgAudio){{ try{{ bgAudio.pause(); bgAudio.src=''; }}catch(_e){{}} bgAudio=null; }}
  if(bgNode)  {{ try{{ bgNode.disconnect(); }}catch(_e){{}} bgNode=null; }}
  if(bgGain)  {{ try{{ bgGain.disconnect(); }}catch(_e){{}} bgGain=null; }}
  if(ctx)     {{ try{{ ctx.close(); }}catch(_e){{}} ctx=null; }}
  if(timer) clearTimeout(timer);
}}

async function start(){{
  if(ctx) return;
  ctx = new (window.AudioContext || window.webkitAudioContext)();

  // --- BUS + LIMITER (volume mais alto e seguro) ---
  bus = ctx.createGain(); bus.gain.value = 1.0;
  limiter = ctx.createDynamicsCompressor();
  limiter.threshold.value = -10; limiter.knee.value = 0; limiter.ratio.value = 20;
  limiter.attack.value = 0.003; limiter.release.value = 0.25;
  bus.connect(limiter).connect(ctx.destination);

  // --- Binaural (L/R) ---
  l = ctx.createOscillator(); r = ctx.createOscillator();
  l.type='sine'; r.type='sine';
  l.frequency.value={fl:.6f}; r.frequency.value={fr:.6f};
  gL = ctx.createGain(); gR = ctx.createGain();
  // ganho do binaural (ajustável no app)
  gL.gain.value = {tg:.4f}; gR.gain.value = {tg:.4f};
  merger = ctx.createChannelMerger(2);
  l.connect(gL).connect(merger,0,0); r.connect(gR).connect(merger,0,1);
  // mistura no BUS (passa pelo limiter)
  merger.connect(bus);
  l.start(); r.start();

  // --- Música de fundo via <audio> ---
  const bg = {bg};
  if (bg) {{
    try {{
      bgAudio = new Audio(bg);
      bgAudio.loop = true;
      bgNode = ctx.createMediaElementSource(bgAudio);
      bgGain = ctx.createGain(); bgGain.gain.value = {g:.4f};

      // Forçar MONO (evita “brigar” com L/R)
      const splitter = ctx.createChannelSplitter(2);
      const mergerMono = ctx.createChannelMerger(2);
      const gA = ctx.createGain(); gA.gain.value = 0.5;
      const gB = ctx.createGain(); gB.gain.value = 0.5;

      bgNode.connect(splitter);
      splitter.connect(gA, 0);
      splitter.connect(gB, 1);
      gA.connect(mergerMono, 0, 0);
      gB.connect(mergerMono, 0, 0);
      mergerMono.connect(bgGain).connect(bus);

      try {{ await bgAudio.play(); }} catch(e) {{ console.warn('Fundo não pôde iniciar:', e); }}
    }} catch(e) {{
      console.warn('Erro no fundo:', e);
    }}
  }}

  timer = setTimeout(()=>stop(), {sec*1000});
}}

function stop(){{
  cleanup();
}}

document.getElementById('bplay').onclick = start;
document.getElementById('bstop').onclick  = stop;
</script>
"""


# -------------------------
# CRUD: patients/intakes/plans/sessions_nova
# -------------------------
def list_patients():
    if BACKEND == "postgres":
        return qall("select id, nome, nascimento from public.patients order by nome asc")
    return sb_select("patients", columns="id,nome,nascimento", order=("nome", True))


def insert_patient(nome: str, telefone: str, email: str, nascimento, notas: str) -> str:
    payload = {
        "nome": nome,
        "telefone": telefone or None,
        "email": email or None,
        "nascimento": str(nascimento) if nascimento else None,
        "notas": notas or None,
    }
    if BACKEND == "postgres":
        row = qone(
            """insert into public.patients (nome, telefone, email, nascimento, notas)
               values (%s,%s,%s,%s,%s) returning id""",
            (nome, telefone or None, email or None, nascimento, notas or None),
        )
        return row["id"]
    row = sb_insert("patients", payload)
    return row["id"]


def insert_intake(
    patient_id: str,
    complaint: str,
    answers: Dict[str, int],
    scores: Dict[str, float],
    flags: Dict[str, bool],
    notes: str,
) -> str:
    payload = {
        "patient_id": patient_id,
        "complaint": complaint or None,
        "answers_json": answers,
        "scores_json": scores,
        "flags_json": flags,
        "notes": notes or None,
    }
    if BACKEND == "postgres":
        row = qone(
            """insert into public.intakes (patient_id, complaint, answers_json, scores_json, flags_json, notes)
               values (%s,%s,%s::jsonb,%s::jsonb,%s::jsonb,%s)
               returning id""",
            (
                patient_id,
                complaint or None,
                json.dumps(answers, ensure_ascii=False),
                json.dumps(scores, ensure_ascii=False),
                json.dumps(flags, ensure_ascii=False),
                notes or None,
            ),
        )
        return row["id"]
    row = sb_insert("intakes", payload)
    return row["id"]


def insert_plan(
    patient_id: str,
    intake_id: str,
    focus: List[Tuple[str, float]],
    selected_names: List[str],
    sessions_qty: int,
    cadence_days: int,
    plan_json: Dict[str, Any],
) -> str:
    payload = {
        "intake_id": intake_id,
        "patient_id": patient_id,
        "focus_json": {"top": focus},
        "selected_protocols": selected_names,
        "sessions_qty": sessions_qty,
        "cadence_days": cadence_days,
        "plan_json": plan_json,
    }
    if BACKEND == "postgres":
        row = qone(
            """insert into public.plans (intake_id, patient_id, focus_json, selected_protocols, sessions_qty, cadence_days, plan_json)
               values (%s,%s,%s::jsonb,%s::jsonb,%s,%s,%s::jsonb)
               returning id""",
            (
                intake_id,
                patient_id,
                json.dumps({"top": focus}, ensure_ascii=False),
                json.dumps(selected_names, ensure_ascii=False),
                sessions_qty,
                cadence_days,
                json.dumps(plan_json, ensure_ascii=False),
            ),
        )
        return row["id"]
    row = sb_insert("plans", payload)
    return row["id"]


def insert_session_nova(plan_id: str, patient_id: str, session_n: int, scheduled_date_str: str, status: str, script: Dict[str, Any]):
    payload = {
        "plan_id": plan_id,
        "patient_id": patient_id,
        "session_n": session_n,
        "scheduled_date": scheduled_date_str,
        "status": status,
        "script_json": script,
    }
    if BACKEND == "postgres":
        qexec(
            """insert into public.sessions_nova (plan_id, patient_id, session_n, scheduled_date, status, script_json)
               values (%s,%s,%s,%s,%s,%s::jsonb)""",
            (plan_id, patient_id, session_n, scheduled_date_str, status, json.dumps(script, ensure_ascii=False)),
        )
    else:
        sb_insert("sessions_nova", payload)



def update_session_nova(session_id: str, updates: Dict[str, Any]):
    """Atualiza campos de sessions_nova (status/notes/scheduled_date/script_json)."""
    if not session_id:
        return
    allowed = {"status", "notes", "scheduled_date", "script_json"}
    upd = {k: v for k, v in (updates or {}).items() if k in allowed}

    if not upd:
        return

    if BACKEND == "postgres":
        sets = []
        params = []
        for k, v in upd.items():
            if k == "script_json":
                sets.append("script_json=%s::jsonb")
                params.append(json.dumps(v, ensure_ascii=False))
            else:
                sets.append(f"{k}=%s")
                params.append(v)
        params.append(session_id)
        qexec(f"update public.sessions_nova set {', '.join(sets)} where id=%s", tuple(params))
    else:
        # supabase
        sb.table("sessions_nova").update(upd).eq("id", session_id).execute()

def delete_session_nova(session_id: str):
    """Apaga uma sessão (mantém as outras)."""
    if not session_id:
        return
    if BACKEND == "postgres":
        qexec("delete from public.sessions_nova where id=%s", (session_id,))
    else:
        sb.table("sessions_nova").delete().eq("id", session_id).execute()

# -------------------------
# HISTÓRICO: intakes / plans (para analisar e reaproveitar anamnese salva)
# -------------------------
def _as_dict(x):
    if x is None:
        return {}
    if isinstance(x, dict):
        return x
    if isinstance(x, str):
        try:
            return json.loads(x)
        except Exception:
            return {}
    return {}

def list_intakes(patient_id: str, limit: int = 30) -> List[Dict[str, Any]]:
    """Lista anamneses do paciente (mais recentes primeiro)."""
    if BACKEND == "postgres":
        return qall(
            """select id, created_at, complaint, answers_json, scores_json, flags_json, notes
                 from public.intakes
                where patient_id=%s
                order by created_at desc
                limit %s""",
            (patient_id, limit),
        )
    return sb_select(
        "intakes",
        columns="id,created_at,complaint,answers_json,scores_json,flags_json,notes",
        eq={"patient_id": patient_id},
        order=("created_at", False),
        limit=limit,
    )

def list_plans(patient_id: str, limit: int = 10) -> List[Dict[str, Any]]:
    """Lista planos gerados do paciente (mais recentes primeiro)."""
    if BACKEND == "postgres":
        return qall(
            """select id, created_at, sessions_qty, cadence_days, selected_protocols, focus_json, plan_json
                 from public.plans
                where patient_id=%s
                order by created_at desc
                limit %s""",
            (patient_id, limit),
        )
    return sb_select(
        "plans",
        columns="id,created_at,sessions_qty,cadence_days,selected_protocols,focus_json,plan_json",
        eq={"patient_id": patient_id},
        order=("created_at", False),
        limit=limit,
    )

def list_sessions_nova(plan_id: str, limit: int = 50) -> List[Dict[str, Any]]:
    if BACKEND == "postgres":
        return qall(
            """select id, session_n, scheduled_date, status, script_json, created_at
                 from public.sessions_nova
                where plan_id=%s
                order by session_n asc
                limit %s""",
            (plan_id, limit),
        )
    return sb_select(
        "sessions_nova",
        columns="id,session_n,scheduled_date,status,script_json,created_at",
        eq={"plan_id": plan_id},
        order=("session_n", True),
        limit=limit,
    )

def apply_intake_to_form(intake_row: Dict[str, Any]):
    """Carrega answers_json para o estado do formulário (compatível com versões antigas e V3)."""
    ans = intake_row.get("answers_json") or {}
    if isinstance(ans, str):
        try:
            ans = json.loads(ans)
        except Exception:
            ans = {}

    # Campos base
    st.session_state[K("att", "complaint")] = intake_row.get("complaint") or ""
    st.session_state[K("att", "notes")] = intake_row.get("notes") or ""

    # --- V3: mapa 0–10 ---
    for s in V3_SYMPTOMS:
        sid = s["id"]
        if K("att", sid) not in st.session_state:
            st.session_state[K("att", sid)] = 0
        # aceita tanto 'sym_*' quanto versões antigas com o mesmo nome
        st.session_state[K("att", sid)] = _clamp_int(ans.get(sid, st.session_state.get(K("att", sid), 0)), 0, 10)

    # --- V3: neuroplasticidade ---
    for q in V3_NEURO:
        qid = q["id"]
        if K("att", qid) not in st.session_state:
            st.session_state[K("att", qid)] = 0
        st.session_state[K("att", qid)] = _clamp_int(ans.get(qid, st.session_state.get(K("att", qid), 0)), 0, 4)

    # tempo (min/dia)
    st.session_state[K("att", "np_time_min")] = _clamp_int(ans.get("np_time_min", st.session_state.get(K("att", "np_time_min"), 5)), 2, 15)

    # --- Compat: questões antigas 0–4 (se existirem no banco) ---
    for q in QUESTIONS:
        qid = q["id"]
        st.session_state[K("att", qid)] = _clamp_int(ans.get(qid, st.session_state.get(K("att", qid), 0)), 0, 4)

    # --- Anamnese física (detalhes) ---
    def _as_list(v):
        if isinstance(v, list):
            return v
        if v in (None, ""):
            return []
        return [v]

    st.session_state[K("att", "phys_dor_local")] = ans.get("phys_dor_local", "") or ""
    st.session_state[K("att", "phys_dor_score")] = _clamp_int(ans.get("phys_dor_score", st.session_state.get(K("att", "phys_dor_score"), 0)), 0, 10)
    st.session_state[K("att", "phys_dor_regioes")] = _as_list(ans.get("phys_dor_regioes"))
    st.session_state[K("att", "phys_hist")] = ans.get("phys_hist", "") or ""
    st.session_state[K("att", "phys_meds_txt")] = ans.get("phys_meds_txt", "") or ""

    # contexto emocional + saúde
    st.session_state[K("att", "phys_emocoes_lida")] = ans.get("phys_emocoes_lida", "Prefiro não responder") or "Prefiro não responder"
    st.session_state[K("att", "phys_emocoes_obs")] = ans.get("phys_emocoes_obs", "") or ""

    st.session_state[K("att", "phys_alergias")] = ans.get("phys_alergias", "Não") or "Não"
    st.session_state[K("att", "phys_alergias_quais")] = ans.get("phys_alergias_quais", "") or ""

    st.session_state[K("att", "phys_cirurgias")] = ans.get("phys_cirurgias", "Não") or "Não"
    st.session_state[K("att", "phys_cirurgias_quais")] = ans.get("phys_cirurgias_quais", "") or ""

    st.session_state[K("att", "phys_hist_familia")] = ans.get("phys_hist_familia", "") or ""

    st.session_state[K("att", "phys_conflito_nivel")] = ans.get("phys_conflito_nivel", "Não") or "Não"
    st.session_state[K("att", "phys_conflito_desc")] = ans.get("phys_conflito_desc", "") or ""

    st.session_state[K("att", "phys_transt_alim")] = ans.get("phys_transt_alim", "Não") or "Não"
    st.session_state[K("att", "phys_transt_alim_desc")] = ans.get("phys_transt_alim_desc", "") or ""

    # --- V4: 9 Origens (compat) ---
    METODOS_O9 = ["Aurímetro", "Pêndulo", "Entrevista", "Outro"]
    metodo = ans.get("o9_metodo") or ans.get("origens9_metodo") or "Aurímetro"
    if metodo not in METODOS_O9:
        metodo = "Outro" if str(metodo).strip() else "Aurímetro"
    st.session_state[K("att", "o9_metodo")] = metodo
    
    # compat: pode vir como dict aninhado em ans["origens9"] ou como chaves planas
    o9 = ans.get("origens9") or ans.get("o9") or {}
    if isinstance(o9, str):
        try:
            o9 = json.loads(o9)
        except Exception:
            o9 = {}
    
    for o in ORIGENS9:
        oid = o["id"]
    
        # medição 0–10
        mval = None
        if isinstance(o9, dict) and isinstance(o9.get(oid), dict):
            mval = o9.get(oid, {}).get("meas")
        if mval is None:
            mval = ans.get(f"o9_{oid}_meas")
        st.session_state[K("att", f"o9_{oid}_meas")] = _clamp_int(mval or 0, 0, 10)
    
        # nota
        nval = None
        if isinstance(o9, dict) and isinstance(o9.get(oid), dict):
            nval = o9.get(oid, {}).get("note")
        if nval is None:
            nval = ans.get(f"o9_{oid}_note")
        st.session_state[K("att", f"o9_{oid}_note")] = str(nval or "")
    
        # evidências (checkbox)
        ev_dict = {}
        if isinstance(o9, dict) and isinstance(o9.get(oid), dict):
            ev_dict = o9.get(oid, {}).get("evidence") or {}
        for ev in (o.get("evidence") or []):
            evid = ev.get("id")
            if not evid:
                continue
            b = None
            if isinstance(ev_dict, dict):
                b = ev_dict.get(evid)
            if b is None:
                b = ans.get(f"o9_{oid}_ev_{evid}")
            st.session_state[K("att", f"o9_{oid}_ev_{evid}")] = bool(b)
    
        # Flags
        flags = intake_row.get("flags_json") or {}
        if isinstance(flags, str):
            try:
                flags = json.loads(flags)
            except Exception:
                flags = {}
        for f in FLAGS:
            fid = f["id"]
            st.session_state[K("att", fid)] = bool(flags.get(fid, st.session_state.get(K("att", fid), False)))
    
    
def reset_att_form_state():
    """Evita 'vazar' estado de um paciente para outro."""
    st.session_state.pop("last_intake_id", None)
    st.session_state[K("att", "complaint")] = ""
    st.session_state[K("att", "notes")] = ""

    # --- V3: mapa 0–10 ---
    for s in V3_SYMPTOMS:
        st.session_state[K("att", s["id"])] = 0

    # --- V3: neuroplasticidade ---
    for q in V3_NEURO:
        st.session_state[K("att", q["id"])] = 0
    st.session_state[K("att", "np_time_min")] = 5

    # --- Anamnese física (detalhes) ---
    st.session_state[K("att", "phys_dor_local")] = ""
    st.session_state[K("att", "phys_dor_score")] = 0
    st.session_state[K("att", "phys_dor_regioes")] = []
    st.session_state[K("att", "phys_hist")] = ""
    st.session_state[K("att", "phys_meds_txt")] = ""

    # contexto emocional + saúde
    st.session_state[K("att", "phys_emocoes_lida")] = "Prefiro não responder"
    st.session_state[K("att", "phys_emocoes_obs")] = ""
    st.session_state[K("att", "phys_alergias")] = "Não"
    st.session_state[K("att", "phys_alergias_quais")] = ""
    st.session_state[K("att", "phys_cirurgias")] = "Não"
    st.session_state[K("att", "phys_cirurgias_quais")] = ""
    st.session_state[K("att", "phys_hist_familia")] = ""
    st.session_state[K("att", "phys_conflito_nivel")] = "Não"
    st.session_state[K("att", "phys_conflito_desc")] = ""
    st.session_state[K("att", "phys_transt_alim")] = "Não"
    st.session_state[K("att", "phys_transt_alim_desc")] = ""

    # Compat: zera sliders antigos 0–4 (se algum trecho ainda ler)
    for q in QUESTIONS:
        st.session_state[K("att", q["id"])] = 0


# --- V4: 9 Origens ---
st.session_state[K("att", "o9_metodo")] = "Aurímetro"
for o in ORIGENS9:
    oid = o["id"]
    st.session_state[K("att", f"o9_{oid}_meas")] = 0
    st.session_state[K("att", f"o9_{oid}_note")] = ""
    for ev in (o.get("evidence") or []):
        evid = ev.get("id")
        if evid:
            st.session_state[K("att", f"o9_{oid}_ev_{evid}")] = False
    # Flags
    for f in FLAGS:
        st.session_state[K("att", f["id"])] = False



def get_frequencies_by_codes(codes: List[str]) -> List[Dict[str, Any]]:
    """Busca detalhes de frequencies pelo code (para mostrar na aba Atendimento)."""
    codes = [str(c).strip().upper() for c in (codes or []) if str(c).strip()]
    if not codes:
        return []
    # dedupe mantendo ordem
    seen = set()
    codes = [c for c in codes if not (c in seen or seen.add(c))]

    if BACKEND == "postgres":
        try:
            rows = qall(
                "select code, nome, hz, tipo, chakra, cor, descricao from public.frequencies where upper(code) = any(%s)",
                (codes,),
            )
            return rows or []
        except Exception:
            # fallback: carrega tudo e filtra
            try:
                all_rows = load_frequencies(None)
                s = set(codes)
                return [r for r in (all_rows or []) if str(r.get("code") or "").strip().upper() in s]
            except Exception:
                return []

    # supabase
    try:
        res = sb.table("frequencies").select("code,nome,hz,tipo,chakra,cor,descricao").in_("code", codes).execute()
        return res.data or []
    except Exception:
        try:
            all_rows = sb_select("frequencies", columns="code,nome,hz,tipo,chakra,cor,descricao", order=("code", True))
            s = set(codes)
            return [r for r in (all_rows or []) if str(r.get("code") or "").strip().upper() in s]
        except Exception:
            return []




# -------------------------
# Impressão (Receituário)
# -------------------------
TEMPLATE_RX_DOCX_DEFAULT = "Receituario_Claudiafito_Template.docx"
import datetime

def get_patient(patient_id: str) -> Optional[Dict[str, Any]]:
    """Busca dados completos do paciente (nome/telefone/email/nascimento/notas)."""
    if not patient_id:
        return None
    if BACKEND == "postgres":
        return qone(
            "select id, nome, telefone, email, nascimento, notas from public.patients where id=%s",
            (patient_id,),
        )
    rows = sb_select(
        "patients",
        columns="id,nome,telefone,email,nascimento,notas",
        eq={"id": patient_id},
        limit=1,
    )
    return rows[0] if rows else None


def _fmt_date_br(x) -> str:
    if not x:
        return ""
    try:
        if isinstance(x, (datetime.date, datetime.datetime)):
            return x.strftime("%d/%m/%Y")
        # strings ISO
        s = str(x)
        # aceita "YYYY-MM-DD" ou "YYYY-MM-DDTHH:MM:SS..."
        d = datetime.date.fromisoformat(s[:10])
        return d.strftime("%d/%m/%Y")
    except Exception:
        return str(x)


def _fmt_time_min_from_seconds(sec: Optional[int]) -> str:
    if sec is None:
        return ""
    try:
        sec = int(sec)
        if sec < 60:
            return f"{sec}s"
        m = sec // 60
        r = sec % 60
        return f"{m} min" + (f" {r}s" if r else "")
    except Exception:
        return str(sec)


_DOMAIN_LABEL = {
    "sono": "Sono",
    "ansiedade": "Ansiedade",
    "humor_baixo": "Humor baixo",
    "exaustao": "Exaustão",
    "pertencimento": "Pertencimento",
    "tensao": "Tensão",
    "ruminacao": "Ruminação",
}

_DOMAIN_OBJ = {
    "sono": "promover relaxamento e higiene do sono",
    "ansiedade": "regular sistema nervoso e reduzir ansiedade/agitação",
    "humor_baixo": "elevar vitalidade e reorganizar energia emocional",
    "exaustao": "repor energia e reduzir sobrecarga",
    "pertencimento": "fortalecer pertencimento, autocompaixão e segurança interna",
    "tensao": "reduzir tensão muscular e estado de alerta",
    "ruminacao": "acalmar mente repetitiva e melhorar foco/presença",
}


# -----------------------------
# Resumo / motivo dos domínios (para orientar o terapeuta)
# -----------------------------
_DOMAIN_RATIONALE: Dict[str, Dict[str, str]] = {
    "sono": {
        "motivo": "Sono é o principal marcador de recuperação (SNC, hormônios do estresse, imunidade).",
        "sinais": "Insônia, despertares, sono não reparador, sonolência diurna.",
        "direcao": "Regular rotina, relaxamento (α/θ), higiene do sono, luz suave.",
    },
    "ansiedade": {
        "motivo": "Ansiedade sinaliza hiperativação (simpático/ruminação) e costuma travar autocuidado.",
        "sinais": "Aperto no peito, inquietação, pensamentos acelerados, tensão muscular.",
        "direcao": "Aterramento, respiração, α/θ, rituais de pausa, apoio emocional.",
    },
    "humor_baixo": {
        "motivo": "Humor baixo indica redução de vitalidade emocional (anhedonia/desânimo) e muda a prioridade do cuidado.",
        "sinais": "Tristeza frequente, perda de prazer, baixa motivação, isolamento, choro fácil.",
        "direcao": "Acolhimento, ativação suave (rotina + pequenos passos), luz/cores quentes moderadas, suporte psicológico se necessário.",
    },
    "exaustao": {
        "motivo": "Exaustão aponta sobrecarga (burnout) e pouca recuperação, afetando sono, foco e imunidade.",
        "sinais": "Cansaço constante, sem energia, irritabilidade, procrastinação, sensação de 'no limite'.",
        "direcao": "Recuperação (sono), reduzir exigências, sessões mais curtas, aterramento, práticas restaurativas e limites.",
    },
    "tensao": {
        "motivo": "Tensão corporal costuma refletir alerta crônico (simpático) e somatização de estresse/ansiedade.",
        "sinais": "Ombros/mandíbula travados, dores musculares, respiração curta, corpo em prontidão.",
        "direcao": "Relaxamento progressivo, respiração, alongamento leve, calor/local, cromoterapia suave e descarrego.",
    },
    "ruminacao": {
        "motivo": "Ruminação mantém o cérebro em looping de ameaça, piora ansiedade e atrapalha presença e sono.",
        "sinais": "Pensamentos repetitivos, dificuldade de foco, mente acelerada, preocupação constante.",
        "direcao": "Treino de atenção (mindfulness), escrita terapêutica, α/θ, reduzir estímulos e organizar rotina.",
    },
    "depressao": {
        "motivo": "Humor rebaixado reduz energia/engajamento e pode exigir cuidado/encaminhamento.",
        "sinais": "Desânimo, apatia, perda de prazer, desesperança, isolamento.",
        "direcao": "Acolhimento, rotina mínima, luz/sons suaves, rede de apoio, encaminhar se necessário.",
    },
    "pertencimento": {
        "motivo": "Pertencimento/segurança social impacta autoestima, limites, decisões e vínculo terapêutico.",
        "sinais": "Solidão, vergonha, sensação de não ter lugar, autojulgamento.",
        "direcao": "Práticas de acolhimento, reconexão, trabalho de limites e apoio comunitário.",
    },
    "estresse": {
        "motivo": "Estresse alto mantém o corpo em alerta e piora sono, dor, digestão e ansiedade.",
        "sinais": "Cansaço, irritabilidade, tensão, sensação de sobrecarga constante.",
        "direcao": "Regulação do sistema nervoso, pausas, respiração, α/θ, redução de estímulos.",
    },
    "energia": {
        "motivo": "Energia baixa aponta exaustão e limita a capacidade de sustentar mudanças.",
        "sinais": "Fadiga, procrastinação, falta de motivação, “sem bateria”.",
        "direcao": "Recuperação (sono), organização do dia, sessões mais curtas e progressivas.",
    },
    "dor": {
        "motivo": "Dor persistente reorganiza o SNC, aumenta estresse e reduz qualidade de vida.",
        "sinais": "Dores recorrentes, travas, tensão, piora com estresse/sono ruim.",
        "direcao": "Relaxamento, analgesia suave, alongamento leve, encaminhar se houver sinais de alarme.",
    },
    "digestao": {
        "motivo": "Eixo intestino-cérebro influencia humor, ansiedade e inflamação.",
        "sinais": "Inchaço, refluxo, intestino preso/solto, desconforto pós-refeição.",
        "direcao": "Rotina alimentar, chás suaves, respiração/relaxamento e observação de gatilhos.",
    },
    "respiracao": {
        "motivo": "Respiração é a alavanca mais rápida para regular ansiedade e tensão.",
        "sinais": "Falta de ar, respiração curta, aperto, crises de ansiedade.",
        "direcao": "Treino respiratório, coerência cardíaca, θ/α e relaxamento de peitoral/diafragma.",
    },
    "imunidade": {
        "motivo": "Imunidade baixa costuma vir com estresse crônico, sono ruim e inflamação.",
        "sinais": "Doenças recorrentes, alergias, fadiga, inflamação frequente.",
        "direcao": "Sono, hidratação, manejo de estresse e apoio suave (sem promessas médicas).",
    },
    "cabeca": {
        "motivo": "Cefaleia reflete tensão, sono ruim, estresse, visão/mandíbula/cervical.",
        "sinais": "Dor de cabeça, enxaqueca, pressão na nuca/testa.",
        "direcao": "Relaxamento, cervical, hidratação, reduzir telas e observar gatilhos.",
    },
    "circulacao": {
        "motivo": "Circulação influencia energia, dor, frio em extremidades e recuperação.",
        "sinais": "Frio em mãos/pés, formigamento, câimbras, inchaço leve.",
        "direcao": "Movimento leve, respiração, hidratação, atenção a sinais de alarme.",
    },
}

def build_domains_summary_df() -> pd.DataFrame:
    """Tabela-guia: domínio → motivo → sinais → direção terapêutica."""
    rows: List[Dict[str, Any]] = []

    # Ordem: DOMAINS (principal) + quaisquer domínios extras que existirem no dicionário de resumo
    ordered = list(DOMAINS)
    for k in (_DOMAIN_RATIONALE or {}).keys():
        if k not in ordered:
            ordered.append(k)

    for d in ordered:
        r = (_DOMAIN_RATIONALE or {}).get(d, {}) or {}
        rows.append(
            {
                "Domínio": _DOMAIN_LABEL.get(d, d),
                "Motivo": r.get("motivo", ""),
                "Sinais comuns quando alto": r.get("sinais", ""),
                "Direção terapêutica (geral)": r.get("direcao", ""),
            }
        )

    return pd.DataFrame(
        rows,
        columns=["Domínio", "Motivo", "Sinais comuns quando alto", "Direção terapêutica (geral)"],
    )

def _join_list(x, sep=", "):
    if not x:
        return ""
    if isinstance(x, str):
        return x
    if isinstance(x, list):
        parts = []
        for it in x:
            if it is None:
                continue
            if isinstance(it, dict):
                # tenta campos comuns
                nome = it.get("nome") or it.get("erva") or it.get("cristal") or it.get("item") or ""
                poso = it.get("posologia") or it.get("dose") or ""
                preparo = it.get("preparo") or it.get("como_usar") or ""
                s = str(nome).strip()
                if preparo:
                    s += f" — {preparo}"
                if poso:
                    s += f" — {poso}"
                if not s.strip():
                    s = json.dumps(it, ensure_ascii=False)
                parts.append(s)
            else:
                parts.append(str(it))
        return sep.join([p for p in parts if p.strip()])
    return str(x)


def _fmt_hz_range(x) -> str:
    """Formata range de Hz vindo do JSON (ex.: [8,12] -> '8–12 Hz')."""
    if x is None:
        return ""
    if isinstance(x, (list, tuple)) and len(x) >= 2:
        a, b = x[0], x[1]
        try:
            a = float(a); b = float(b)
            if abs(a - b) < 1e-9:
                return f"{a:.1f} Hz".replace('.0', '')
            return f"{a:.1f}–{b:.1f} Hz".replace('.0', '')
        except Exception:
            return f"{a}–{b} Hz"
    try:
        v = float(x)
        return f"{v:.1f} Hz".replace('.0', '')
    except Exception:
        return str(x)


def _fmt_minutes_from_any(d: dict) -> str:
    """Extrai duração em minutos de campos comuns (dur_min, duracao_min, duracao_s...)."""
    for k in ["dur_min", "duracao_min", "min"]:
        if k in d and d.get(k) is not None:
            try:
                return f"{float(d.get(k)):.0f} min"
            except Exception:
                return str(d.get(k))
    for k in ["duracao_s", "dur_s", "seconds"]:
        if k in d and d.get(k) is not None:
            try:
                sec = int(float(d.get(k)))
                if sec < 60:
                    return f"{sec}s"
                m = sec // 60
                r = sec % 60
                return f"{m} min" + (f" {r}s" if r else "")
            except Exception:
                return str(d.get(k))
    return ""


def _format_binaural_protocolos(rows: Any) -> str:
    """Transforma a sugestão binaural dos protocolos (lista de dict) em texto legível."""
    if not rows:
        return ""
    if isinstance(rows, dict):
        rows = [rows]
    if isinstance(rows, str):
        # já veio formatado
        return rows.strip()

    out_lines: List[str] = []
    if isinstance(rows, list):
        for r in rows:
            if not isinstance(r, dict):
                s = str(r).strip()
                if s:
                    out_lines.append(f"• {s}")
                continue

            prot = str(r.get("protocolo") or r.get("card") or "").strip()
            mode = str(r.get("mode") or r.get("faixa") or "").strip()
            beat = _fmt_hz_range(r.get("beat_hz") or r.get("beat") or "")
            dur = _fmt_minutes_from_any(r)
            obs = str(r.get("obs") or r.get("nota") or r.get("notes") or "").strip()

            parts = []
            if prot:
                parts.append(prot)
            if mode:
                parts.append(mode)
            if beat:
                parts.append(f"Beat {beat}")
            if dur:
                parts.append(dur)

            line = " — ".join([p for p in parts if p])
            if obs:
                line = (line + f". Obs: {obs}") if line else obs

            if line.strip():
                out_lines.append(f"• {line}")

    # Limita para não poluir o receituário
    return "\n".join(out_lines[:8]).strip()


def _summarize_cama_rows(cama_rows: Any) -> str:
    """Formata sugestão de Cama de Cristal (JSON) em texto amigável."""
    if not cama_rows:
        return ""
    if isinstance(cama_rows, str):
        return cama_rows.strip()

    # Normaliza em lista
    rows = cama_rows
    if isinstance(rows, dict):
        rows = [rows]

    # Agrupa por protocolo quando possível
    grouped: Dict[str, List[dict]] = {}
    loose: List[dict] = []

    if isinstance(rows, list):
        for it in rows:
            if isinstance(it, dict):
                prot = str(it.get("protocolo") or it.get("card") or "").strip()
                if prot:
                    grouped.setdefault(prot, []).append(it)
                else:
                    loose.append(it)

    def _fmt_step(step: dict, idx: int) -> str:
        chakra = step.get("chakra") or step.get("Chakra") or ""
        cor = step.get("cor") or step.get("Cor") or ""
        mins = step.get("min") or step.get("mins") or step.get("duracao_min") or step.get("tempo") or step.get("Tempo") or ""
        bits = [str(chakra).strip(), str(cor).strip()]
        bits = [b for b in bits if b]
        base = " — ".join(bits) if bits else ""
        if mins != "":
            try:
                base = (base + f" — {float(mins):.0f} min") if base else f"{float(mins):.0f} min"
            except Exception:
                base = (base + f" — {mins}") if base else str(mins)
        if not base:
            base = json.dumps(step, ensure_ascii=False)
        return f"{idx}. {base}"

    lines: List[str] = []

    # Protocolos (ordenados)
    for prot in sorted(grouped.keys()):
        items = grouped[prot]
        # Caso 1: dict com "sequencia" (lista de passos)
        seq = None
        for r in items:
            if isinstance(r.get("sequencia"), list):
                seq = r.get("sequencia")
                obs = str(r.get("obs") or r.get("nota") or "").strip()
                lines.append(f"{prot}:")
                for i, step in enumerate(seq, start=1):
                    if isinstance(step, dict):
                        lines.append("  " + _fmt_step(step, i))
                    else:
                        lines.append(f"  {i}. {step}")
                if obs:
                    lines.append(f"  Obs: {obs}")
                lines.append("")  # linha em branco
                break

        if seq is not None:
            continue

        # Caso 2: itens já vêm como passos (um por linha)
        step_like = [r for r in items if isinstance(r, dict) and (r.get("chakra") or r.get("cor") or r.get("min") or r.get("duracao_min"))]
        if step_like:
            lines.append(f"{prot}:")
            for i, step in enumerate(step_like, start=1):
                lines.append("  " + _fmt_step(step, i))
            lines.append("")
            continue

        # Caso 3: fallback (qualquer dict)
        lines.append(f"{prot}: {json.dumps(items[0], ensure_ascii=False)}")
        lines.append("")
    # Loose items
    for r in loose[:6]:
        try:
            lines.append(json.dumps(r, ensure_ascii=False))
        except Exception:
            lines.append(str(r))

    # limpa linhas finais vazias
    while lines and not lines[-1].strip():
        lines.pop()

    return "\n".join(lines).strip()

def _build_receituario_data_from_plan(patient: Dict[str, Any], plan_row: Dict[str, Any], sessions: List[Dict[str, Any]]) -> Dict[str, Any]:
    plan_json = plan_row.get("plan_json") or {}
    complaint = (plan_json.get("complaint") or "") if isinstance(plan_json, dict) else ""
    plan_date = (plan_json.get("date") or "") if isinstance(plan_json, dict) else ""
    scores = (plan_json.get("scores") or {}) if isinstance(plan_json, dict) else {}
    focus = (plan_json.get("focus") or []) if isinstance(plan_json, dict) else []
    selected_protocols = plan_row.get("selected_protocols") or plan_json.get("selected_protocols") or plan_row.get("selected_protocols_json") or plan_row.get("selected_protocols") or []
    merged_plan = (plan_json.get("plan") or {}) if isinstance(plan_json, dict) else {}

    # Frequências auxiliares (codes)
    freqs = plan_json.get("frequencias") if isinstance(plan_json, dict) else None
    freq_codes = []
    if isinstance(freqs, list):
        for f in freqs:
            if isinstance(f, dict) and f.get("code"):
                freq_codes.append(str(f["code"]))
            elif isinstance(f, str):
                freq_codes.append(f)
    freq_codes = [c for c in freq_codes if str(c).strip()]

    # Áudio
    audio = (plan_json.get("audio") or {}) if isinstance(plan_json, dict) else {}
    binaural = (audio.get("binaural") or {}) if isinstance(audio, dict) else {}

    # Cama de cristal: pega do primeiro script_json se existir
    cama_txt = ""
    if sessions:
        sj = sessions[0].get("script_json") or {}
        cama_txt = _summarize_cama_rows(sj.get("cama_cristal_sugestao"))
    if not cama_txt:
        # fallback: chakras/cor do plano consolidado
        chakras = merged_plan.get("chakras_prioritarios") if isinstance(merged_plan, dict) else None
        if chakras:
            cama_txt = "Chakras prioritários: " + _join_list(chakras)
    # Binaural dos protocolos (sugestão)
    binaural_protocolos_txt = ""
    if sessions:
        sj = sessions[0].get("script_json") or {}
        binaural_protocolos_txt = _format_binaural_protocolos(sj.get("binaural_protocolos_sugestao"))

    # Fito / cristais / alertas
    fito = merged_plan.get("fito_sugerida") if isinstance(merged_plan, dict) else None
    cristais = merged_plan.get("cristais_sugeridos") if isinstance(merged_plan, dict) else None
    alertas = merged_plan.get("alertas") if isinstance(merged_plan, dict) else None

    fito_txt = _join_list(fito, sep="\n• ")
    if fito_txt:
        fito_txt = "• " + fito_txt if "\n" in fito_txt else fito_txt

    cristais_txt = _join_list(cristais)
    cuidados = ""
    if alertas:
        cuidados = _join_list(alertas, sep="\n• ")
        cuidados = "• " + cuidados if "\n" in cuidados else cuidados
    # Observações de segurança (padrão)
    cuidados_base = (
        "• Use fones e volume moderado.\n"
        "• Interrompa se houver desconforto (tontura, dor de cabeça, náusea).\n"
        "• Terapia integrativa não substitui acompanhamento médico/psicológico."
    )
    cuidados = (cuidados.strip() + "\n" if cuidados else "") + cuidados_base

    # Objetivos (derivados do foco)
    objetivos = []
    if isinstance(focus, list) and focus:
        for d, sc in focus[:3]:
            objetivos.append(f"• {_DOMAIN_LABEL.get(d, d)}: {_DOMAIN_OBJ.get(d, 'equilibrar este domínio')} (score {sc:.1f}%)")
    objetivos_txt = "\n".join(objetivos) if objetivos else "• Acolher queixa principal e promover regulação."

    # Sessões: tabela
    sess_rows = []
    for s in sessions[:8]:
        sess_rows.append({
            "n": s.get("session_n"),
            "data": _fmt_date_br(s.get("scheduled_date") or (s.get("script_json") or {}).get("scheduled_date")),
            "status": s.get("status") or (s.get("script_json") or {}).get("status") or "",
        })

    # Sugestão de sessões (texto)
    qty = plan_row.get("sessions_qty") or plan_json.get("sessions_qty")
    cadence = plan_row.get("cadence_days") or plan_json.get("cadence_days")
    sessoes_txt = ""
    if qty and cadence:
        semanas = max(1, int(math.ceil((int(qty) * int(cadence)) / 7)))
        sessoes_txt = f"{int(qty)} sessões em ~{semanas} semanas (a cada {int(cadence)} dias)."

    # Binaural (texto)
    carrier = binaural.get("carrier_hz")
    beat = binaural.get("beat_hz")
    dur = binaural.get("duracao_s")
    binaural_txt = ""
    if carrier is not None and beat is not None:
        binaural_txt = f"Carrier {float(carrier):.0f} Hz | Beat {float(beat):.1f} Hz | Duração {_fmt_time_min_from_seconds(dur)}"
        if binaural_protocolos_txt:
            binaural_txt += "\n\nSugestões por protocolo:\n" + binaural_protocolos_txt

    freq_aux_txt = ", ".join(freq_codes) if freq_codes else ""

    return {
        "patient_nome": patient.get("nome") or "",
        "patient_nasc": _fmt_date_br(patient.get("nascimento")),
        "patient_whats": patient.get("telefone") or "",
        "patient_email": patient.get("email") or "",
        "sessao_data": _fmt_date_br(plan_date) if plan_date else _fmt_date_br(datetime.date.today()),
        "queixa": complaint or "",
        "scores": scores or {},
        "focus": focus or [],
        "protocolos": selected_protocols or [],
        "objetivos_txt": objetivos_txt,
        "sessoes_txt": sessoes_txt,
        "binaural_txt": binaural_txt,
        "freq_aux_txt": freq_aux_txt,
        "cama_txt": cama_txt,
        "fito_txt": fito_txt or "",
        "cristais_txt": cristais_txt or "",
        "cuidados_txt": cuidados,
        "sess_rows": sess_rows,
    }


def _docx_replace_in_paragraph(paragraph, mapping: Dict[str, str]):
    # substitui mantendo o estilo do primeiro run
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)
    if new != full:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""


def _docx_replace_everywhere(doc, mapping: Dict[str, str]):
    for p in doc.paragraphs:
        _docx_replace_in_paragraph(p, mapping)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _docx_replace_in_paragraph(p, mapping)


def generate_receituario_docx_bytes(data: Dict[str, Any], template_file: Optional[io.BytesIO] = None) -> bytes:
    """Gera DOCX preenchido a partir do template."""
    from docx import Document  # lazy import

    # Carrega template
    if template_file is not None:
        doc = Document(template_file)
    else:
        base_dir = os.path.dirname(__file__) if "__file__" in globals() else "."
        path = os.path.join(base_dir, TEMPLATE_RX_DOCX_DEFAULT)
        doc = Document(path)

    # 1) Preenchimentos diretos
    mapping = {
        "[NOME COMPLETO]": data.get("patient_nome", ""),
        "[DD/MM/AAAA]": data.get("patient_nasc", ""),
        "[WhatsApp]": data.get("patient_whats", ""),
        "[E-mail]": data.get("patient_email", ""),
        "[DATA DA SESSÃO]": data.get("sessao_data", ""),
        "[QUEIXA PRINCIPAL]": data.get("queixa", ""),
        "[FOCO 1]": "",
        "[FOCO 2]": "",
        "[FOCO 3]": "",
        "[Ex.: regular sistema nervoso autônomo (ansiedade, ruminação, tensão).]": data.get("objetivos_txt", ""),
        "[Ex.: 6 sessões em 6–8 semanas (1/semana).]": data.get("sessoes_txt", ""),
        "[Ex.: Theta 6 Hz 15 min → Alpha 10 Hz 10 min.]": data.get("binaural_txt", ""),
        "[Cama de Cristal]": data.get("cama_txt", ""),
        "[Fitoenergética / ervas]": data.get("fito_txt", ""),
        "[Cristais sugeridos]": data.get("cristais_txt", ""),
        "[Cuidados]": data.get("cuidados_txt", ""),
    }

    # Focos (top 3)
    focus = data.get("focus") or []
    for i in range(3):
        if i < len(focus):
            d, sc = focus[i]
            mapping[f"[FOCO {i+1}]"] = f"{_DOMAIN_LABEL.get(d, d)} ({float(sc):.1f}%)"
        else:
            mapping[f"[FOCO {i+1}]"] = ""

    _docx_replace_everywhere(doc, mapping)

    # 2) Scores: tabela com [__]
    scores = data.get("scores") or {}
    try:
        # tabela de domínios é a 3ª (índice 2) no template
        t = doc.tables[2]
        # linhas 1..7
        for ri in range(1, min(len(t.rows), 8)):
            dom = list(_DOMAIN_LABEL.keys())[ri-1]
            val = scores.get(dom, "")
            # coluna 1
            cell = t.rows[ri].cells[1]
            # substitui qualquer [__] que existir
            for p in cell.paragraphs:
                _docx_replace_in_paragraph(p, {"[__]": f"{val}%" if val != "" else ""})
    except Exception:
        pass

    # 3) Binaural / Frequências auxiliares: tabela 4 (índice 4)
    try:
        t = doc.tables[4]
        # expected rows: Carrier, Beat, Duração, Frequências auxiliares
        # As células de valor têm [___]
        # Vamos preencher pelo texto da primeira coluna
        binaural_txt = data.get("binaural_txt", "")
        # tentar extrair carrier/beat/dur do texto, mas se não, preencher tudo no primeiro
        carrier_val = ""
        beat_val = ""
        dur_val = ""
        if binaural_txt:
            # formatos: Carrier 220 Hz | Beat 10.0 Hz | Duração 2 min
            m = re.search(r"Carrier\s+([0-9.]+)\s*Hz", binaural_txt)
            if m: carrier_val = f"{float(m.group(1)):.0f} Hz"
            m = re.search(r"Beat\s+([0-9.]+)\s*Hz", binaural_txt)
            if m: beat_val = f"{float(m.group(1)):.1f} Hz"
            m = re.search(r"Duração\s+(.+?)(\n|$)", binaural_txt)
            if m: dur_val = m.group(1).strip()
        freq_aux = data.get("freq_aux_txt", "")

        for row in t.rows[1:]:
            label = (row.cells[0].text or "").lower()
            cell = row.cells[1]
            repl = {}
            if "carrier" in label:
                repl["[___]"] = carrier_val
            elif "beat" in label or "batida" in label:
                repl["[___]"] = beat_val
            elif "duração" in label or "duracao" in label:
                repl["[___]"] = dur_val
            elif "auxiliares" in label:
                repl["[___]"] = freq_aux
            if repl:
                for p in cell.paragraphs:
                    _docx_replace_in_paragraph(p, repl)
    except Exception:
        pass

    # 4) Sessões: tabela 5 (índice 5) — até 8 linhas
    try:
        t = doc.tables[5]
        rows = data.get("sess_rows") or []
        for i in range(1, min(len(t.rows), 9)):
            # i-1 é índice em rows
            if i-1 < len(rows):
                r = rows[i-1]
                n = str(r.get("n") or i)
                d = str(r.get("data") or "")
                s = str(r.get("status") or "")
            else:
                n, d, s = "", "", ""
            # col0 [1], col1 [DATA], col2 [Status]
            for p in t.rows[i].cells[0].paragraphs:
                _docx_replace_in_paragraph(p, {f"[{i}]": n})
            for p in t.rows[i].cells[1].paragraphs:
                _docx_replace_in_paragraph(p, {"[DATA]": d})
            for p in t.rows[i].cells[2].paragraphs:
                _docx_replace_in_paragraph(p, {"[Status]": s})
    except Exception:
        pass

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_receituario_pdf_bytes(data: Dict[str, Any]) -> bytes:
    """Gera um PDF simples (A4) com as mesmas informações do receituário."""
    if not HAS_REPORTLAB:
        raise RuntimeError("PDF indisponível: dependência 'reportlab' não está instalada. Baixe o DOCX (imprimível) ou adicione 'reportlab' ao requirements.txt.")

    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Receituário / Orientações de Atendimento", styles["Title"]))
    story.append(Spacer(1, 8))

    patient_line = f"<b>Paciente:</b> {data.get('patient_nome','')} &nbsp;&nbsp; <b>Nasc.:</b> {data.get('patient_nasc','')}"
    contact_line = f"<b>WhatsApp:</b> {data.get('patient_whats','')} &nbsp;&nbsp; <b>E-mail:</b> {data.get('patient_email','')}"
    story.append(Paragraph(patient_line, styles["Normal"]))
    story.append(Paragraph(contact_line, styles["Normal"]))
    story.append(Paragraph(f"<b>Data:</b> {data.get('sessao_data','')}", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph(f"<b>Queixa principal:</b> {data.get('queixa','')}", styles["Normal"]))
    story.append(Spacer(1, 8))

    # Scores
    scores = data.get("scores") or {}
    score_rows = [["Domínio", "Score"]]
    for dom in DOMAINS:
        score_rows.append([_DOMAIN_LABEL.get(dom, dom), f"{scores.get(dom,'')}%"])
    tbl = Table(score_rows, hAlign="LEFT", colWidths=[8*cm, 3*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(Paragraph("<b>Pontuações (anamnese):</b>", styles["Normal"]))
    story.append(tbl)
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>Objetivos terapêuticos:</b>", styles["Normal"]))
    story.append(Paragraph(data.get("objetivos_txt","").replace("\n","<br/>"), styles["Normal"]))
    story.append(Spacer(1, 8))

    if data.get("sessoes_txt"):
        story.append(Paragraph(f"<b>Plano de sessões:</b> {data.get('sessoes_txt','')}", styles["Normal"]))
        story.append(Spacer(1, 8))

    if data.get("binaural_txt"):
        story.append(Paragraph("<b>Binaural (em casa ou na sessão):</b>", styles["Normal"]))
        story.append(Paragraph(data.get("binaural_txt","").replace("\n","<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))

    if data.get("freq_aux_txt"):
        story.append(Paragraph(f"<b>Frequências auxiliares (codes):</b> {data.get('freq_aux_txt','')}", styles["Normal"]))
        story.append(Spacer(1, 6))

    if data.get("cama_txt"):
        story.append(Paragraph("<b>Cama de cristal (sugestão):</b>", styles["Normal"]))
        story.append(Paragraph(data.get("cama_txt", "").replace("\n","<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))

    if data.get("cristais_txt"):
        story.append(Paragraph(f"<b>Cristais sugeridos:</b> {data.get('cristais_txt','')}", styles["Normal"]))
        story.append(Spacer(1, 6))

    if data.get("fito_txt"):
        story.append(Paragraph("<b>Fitoenergética / ervas (orientação):</b>", styles["Normal"]))
        story.append(Paragraph(data.get("fito_txt","").replace("\n","<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Cuidados:</b>", styles["Normal"]))
    story.append(Paragraph(data.get("cuidados_txt","").replace("\n","<br/>"), styles["Normal"]))
    story.append(Spacer(1, 10))

    # Sessões
    sess = data.get("sess_rows") or []
    if sess:
        sess_tbl = [["#", "Data", "Status"]] + [[str(r.get("n") or ""), r.get("data") or "", r.get("status") or ""] for r in sess]
        t = Table(sess_tbl, hAlign="LEFT", colWidths=[1*cm, 4*cm, 6*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ]))
        story.append(Paragraph("<b>Cronograma de sessões (planejado):</b>", styles["Normal"]))
        story.append(t)

    doc.build(story)
    return buf.getvalue()


# -------------------------
# UI helpers
# -------------------------
def K(*parts: str) -> str:
    return "__".join(parts)


# --- Anamnese física (detalhes) ---
PHYS_DOR_REGIOES = [
    "Cabeça / enxaqueca",
    "Pescoço",
    "Ombros",
    "Coluna cervical",
    "Coluna torácica",
    "Coluna lombar",
    "Quadril",
    "Joelhos",
    "Pés / tornozelos",
    "Abdômen",
    "Outros",
]



def _json_str(x: Any) -> str:
    """String segura para mostrar em grid (mantém dict/list como JSON)."""
    if x is None:
        return ""
    if isinstance(x, (dict, list)):
        try:
            return json.dumps(x, ensure_ascii=False)
        except Exception:
            return str(x)
    return str(x)

def json_to_df(obj: Any, name: str = "item") -> pd.DataFrame:
    """
    Converte dict/list/valor em DataFrame para visualização.
    - dict -> colunas: chave, valor
    - list[dict] -> normaliza colunas
    - list[scalar] -> 1 coluna (name)
    - scalar -> 1 coluna (name)
    """
    if obj is None:
        return pd.DataFrame(columns=[name])
    # tenta parsear string JSON
    if isinstance(obj, str):
        s = obj.strip()
        if (s.startswith("{") and s.endswith("}")) or (s.startswith("[") and s.endswith("]")):
            try:
                obj = json.loads(s)
            except Exception:
                return pd.DataFrame([{name: obj}])
        else:
            return pd.DataFrame([{name: obj}])

    if isinstance(obj, dict):
        rows = [{"chave": k, "valor": _json_str(v)} for k, v in obj.items()]
        return pd.DataFrame(rows)

    if isinstance(obj, list):
        if not obj:
            return pd.DataFrame(columns=[name])
        if all(isinstance(it, dict) for it in obj):
            # normaliza e garante que valores complexos virem string
            rows = []
            for it in obj:
                row = {k: _json_str(v) for k, v in it.items()}
                rows.append(row)
            return pd.DataFrame(rows)
        # lista de valores
        return pd.DataFrame([{name: _json_str(it)} for it in obj])

    return pd.DataFrame([{name: _json_str(obj)}])

# ---- CHAVES únicas para widgets binaurais (não duplicar estado) ----
KEY_CARRIER = K("binaural", "carrier")
KEY_BEAT    = K("binaural", "beat")
KEY_DUR_S   = K("binaural", "dur_s")
KEY_BG_GAIN = K("binaural", "bg_gain")
KEY_TONE_GAIN = K("binaural", "tone_gain")

st.title("🌿✨ SISTEMA INTEGRADO DE GERAÇÃO DE TERAPIAS INTEGRATIVAS - DOCE CONEXÃO ✨🌿")
st.caption("Geração de Planos Terapêuticos integrados utilizando terapias: Fitoterapia, Frequnciais Solfeggios e Binaurais, Cromoterapia e Cristaloterapia.")

tabs = st.tabs(["Atendimento", "Binaural"])

# Shared binaural settings (1 fonte de verdade = os widgets)
st.session_state.setdefault(KEY_CARRIER, 220.0)
st.session_state.setdefault(KEY_BEAT, 10.0)
st.session_state.setdefault(KEY_DUR_S, 120)     # duração em SEGUNDOS (igual no app antigo)
st.session_state.setdefault(KEY_BG_GAIN, 0.12)
st.session_state.setdefault(KEY_TONE_GAIN, 0.30)  # volume do binaural (WebAudio)
st.session_state.setdefault("extra_freq_codes", [])

# Também expõe em chaves "antigas" para o Atendimento ler (compatibilidade)
st.session_state.setdefault("binaural_carrier", float(st.session_state[KEY_CARRIER]))
st.session_state.setdefault("binaural_beat", float(st.session_state[KEY_BEAT]))
st.session_state.setdefault("binaural_dur_s", int(st.session_state[KEY_DUR_S]))
st.session_state.setdefault("binaural_bg_gain", float(st.session_state[KEY_BG_GAIN]))
st.session_state.setdefault("binaural_tone_gain", float(st.session_state[KEY_TONE_GAIN]))

# -------------------------
# TAB: BINAURAL
# -------------------------
with tabs[1]:
    st.subheader("Binaural")

    band_map = {
        "Delta (1–4 Hz)": 3.0,
        "Theta (4–8 Hz)": 6.0,
        "Alpha (8–12 Hz)": 10.0,
        "Beta baixa (12–18 Hz)": 15.0,
        "Gamma (30–45 Hz)": 40.0,
    }
    bcol1, bcol2 = st.columns([2, 1])
    faixa = bcol1.selectbox("Faixa de ondas (atalho)", list(band_map.keys()), index=2, key=K("binaural", "band"))
    if bcol2.button("Aplicar faixa", key=K("binaural", "apply_band")):
        # A faixa troca SOMENTE o beat (carrier fica como está) — isso é normal.
        st.session_state[KEY_BEAT] = float(band_map[faixa])
        st.session_state["binaural_beat"] = float(st.session_state[KEY_BEAT])
        st.success(f"Batida ajustada para {band_map[faixa]} Hz")
        st.rerun()

    # Presets do banco
    try:
        presets = load_binaural_presets()
    except Exception as e:
        presets = []
        st.warning(f"Não consegui ler binaural_presets: {e}")

    mapa_pres = {p["nome"]: p for p in (presets or []) if p.get("nome")}
    cols_top = st.columns([2, 1])
    preset_names = list(mapa_pres.keys()) or ["(nenhum)"]
    preset_escolhido = cols_top[0].selectbox("Tratamento pré-definido", preset_names, key=K("binaural", "preset"))

    if cols_top[1].button("Aplicar preset", key=K("binaural", "apply_preset")) and preset_escolhido in mapa_pres:
        p = mapa_pres[preset_escolhido]
        # >>> IMPORTANTE: atualizar os mesmos KEYS dos widgets <<<
        st.session_state[KEY_CARRIER] = float(p.get("carrier_hz") or 220.0)
        st.session_state[KEY_BEAT]    = float(p.get("beat_hz") or 10.0)
        dur_min = p.get("duracao_min")
        if dur_min is not None:
            st.session_state[KEY_DUR_S] = int(float(dur_min) * 60)

        # espelha para o Atendimento
        st.session_state["binaural_carrier"] = float(st.session_state[KEY_CARRIER])
        st.session_state["binaural_beat"] = float(st.session_state[KEY_BEAT])
        st.session_state["binaural_dur_s"] = int(st.session_state[KEY_DUR_S])

        st.success("Preset aplicado.")
        st.rerun()

    c1, c2, c3 = st.columns(3)
    carrier = c1.number_input("Carrier (Hz)", 50.0, 1000.0, step=1.0, key=KEY_CARRIER)
    beat    = c2.number_input("Batida (Hz)", 0.5, 45.0, step=0.5, key=KEY_BEAT)
    dur_s   = int(c3.number_input("Duração (s)", 10, 3600, step=5, key=KEY_DUR_S))

    # espelha para o Atendimento
    st.session_state["binaural_carrier"] = float(carrier)
    st.session_state["binaural_beat"] = float(beat)
    st.session_state["binaural_dur_s"] = int(dur_s)

    bt = abs(float(beat))
    fL = max(20.0, float(carrier) - bt / 2.0)
    fR = float(carrier) + bt / 2.0
    mL, mR = st.columns(2)
    mL.metric("Esquerdo (L)", f"{fL:.2f} Hz")
    mR.metric("Direito (R)", f"{fR:.2f} Hz")

    with st.expander("Como funciona?"):
        st.markdown(
            """
**Binaural** = duas frequências **puras** diferentes em cada ouvido → o cérebro percebe a **diferença** como um tom de batida (**beat**).  
**Cálculo:** `L = carrier − beat/2` e `R = carrier + beat/2`.  
Ex.: carrier 220 Hz e beat 10 Hz ⇒ L = **215 Hz**, R = **225 Hz** ⇒ o cérebro tende a sincronizar em **~10 Hz**.
"""
        )

    st.markdown("🎵 Música de fundo (opcional) — do seu computador (como antes)")
    # Volume do binaural (separado do volume do fundo)
    tone_gain = st.slider(
        "Volume do binaural",
        min_value=0.02,
        max_value=0.80,
        step=0.01,
        key=KEY_TONE_GAIN,
        help="Aumente se estiver baixo. Use fones e mantenha volume moderado.",
    )
    # espelha para o Atendimento (compatibilidade)
    st.session_state["binaural_tone_gain"] = float(tone_gain)


    bg_up = st.file_uploader("MP3/WAV/OGG (até 12MB)", type=["mp3", "wav", "ogg"], key=K("binaural", "bg_file"))
    bg_gain = st.slider("Volume do fundo", min_value=0.0, max_value=0.60, step=0.01, key=KEY_BG_GAIN)

    st.session_state["binaural_bg_gain"] = float(bg_gain)

    raw = None
    filename = None
    if bg_up:
        raw = bg_up.read()
        filename = bg_up.name
        st.audio(raw)  # prévia

    bg_url, _mime, err = bytes_to_data_url_safe(raw, filename) if raw else (None, None, None)
    if err:
        st.warning(f"⚠️ {err}")

    st.markdown("▶️ **Player (Tocar/Parar)** — binaural + fundo")
    components.html(
        webaudio_binaural_html(
            float(carrier),
            float(beat),
            seconds=int(dur_s),
            bg_data_url=bg_url,
            bg_gain=float(bg_gain),
            binaural_gain=float(tone_gain),
        ),
        height=140,
    )

    st.markdown("🔔 Frequências auxiliares (Solfeggio + Chakras)")
    try:
        sol = load_frequencies("solfeggio")
        chak = load_frequencies("chakra")
    except Exception as e:
        sol, chak = [], []
        st.warning(f"Não consegui ler frequencies: {e}")

    def _opt(r):
        code = str(r.get("code") or "").strip()
        nome = str(r.get("nome") or "").strip()
        hz = r.get("hz")
        hz_s = f"{float(hz):.2f} Hz" if hz is not None else "—"
        base = code if code else "(sem code)"
        if nome:
            base += f" — {nome}"
        return f"{base} • {hz_s}"

    sol_opts = [_opt(r) for r in sol]
    chak_opts = [_opt(r) for r in chak]
    sol_map = {sol_opts[i]: sol[i].get("code") for i in range(len(sol_opts))}
    chak_map = {chak_opts[i]: chak[i].get("code") for i in range(len(chak_opts))}

    c4, c5 = st.columns(2)
    sel_sol = c4.multiselect("Solfeggio", sol_opts, default=[], key=K("binaural", "sol"))
    sel_chak = c5.multiselect("Chakras", chak_opts, default=[], key=K("binaural", "chak"))

    extra_codes = [sol_map[x] for x in sel_sol if sol_map.get(x)] + [chak_map[x] for x in sel_chak if chak_map.get(x)]
    custom_code = st.text_input("Custom code (opcional)", value="", key=K("binaural", "custom_code"))
    if custom_code.strip():
        extra_codes.append(custom_code.strip().upper())

    seen = set()
    extra_codes = [c for c in extra_codes if not (c in seen or seen.add(c))]
    st.session_state["extra_freq_codes"] = extra_codes

    # WAV de preview/download (20s): usa um ganho proporcional ao "Volume do binaural"
    wav_amp = min(0.95, max(0.05, float(tone_gain) * 4.0))
    wav = synth_binaural_wav(float(carrier), float(beat), seconds=20, sr=44100, amp=float(wav_amp))
    st.audio(wav, format="audio/wav")
    st.download_button(
        "Baixar WAV (20s)",
        data=wav,
        file_name=f"binaural_{int(carrier)}_{beat:.1f}.wav",
        mime="audio/wav",
        key=K("binaural", "dl_wav"),
    )

    with st.expander("Sugestões rápidas por objetivo"):
        st.markdown(
            """
- **Relaxar/ansiedade** → **Theta 5–6 Hz** (15–20 min) e fechar em **Alpha 10 Hz** (5–10 min).  
- **Sono** → **Delta 2–3 Hz** (10–20 min) → **Theta 5–6 Hz** (10–15 min).  
- **Foco calmo** → **Alpha 10 Hz** (10–15 min).  
- **Gamma 40 Hz** → estimulação breve (5–12 min), volume baixo.  
            """
        )

# -------------------------
# TAB: ATENDIMENTO
# -------------------------
with tabs[0]:
    st.markdown("## Atendimento • Diagnóstico Integrativo (9 Origens)")
    st.caption(
        "Fluxo profissional: **medição + evidências + contexto** → gera causas prováveis (9 origens) e traduz para um plano terapêutico (protocolos + sessões). "
        "Obs.: é um **modelo de triagem terapêutica** (não é diagnóstico médico)."
    )

    # -------------------------
    # Sidebar — Paciente
    # -------------------------
    with st.sidebar:
        st.header("Paciente")
        try:
            patients = list_patients()
        except Exception as e:
            patients = []
            st.error(f"Erro ao carregar patients: {e}")

        def lab(p):
            nasc = p.get("nascimento")
            tail = str(p.get("id") or "")[-4:]
            return f'{p.get("nome","(sem nome)")} — {nasc or "s/n"} — {tail}'

        labels = ["— Novo paciente —"] + [lab(p) for p in patients]
        sel = st.selectbox("Selecionar", labels, index=0, key=K("pat", "sel"))

        if sel == "— Novo paciente —":
            nome = st.text_input("Nome", key=K("pat", "nome"))
            telefone = st.text_input("Telefone (opcional)", key=K("pat", "tel"))
            email = st.text_input("E-mail (opcional)", key=K("pat", "email"))
            nascimento = st.date_input("Nascimento (opcional)", value=None, key=K("pat", "nasc"))
            pnotas = st.text_area("Notas (opcional)", key=K("pat", "notas"))
            if st.button("Criar paciente", type="primary", use_container_width=True, key=K("pat", "create")):
                if not nome.strip():
                    st.warning("Informe o nome.")
                else:
                    try:
                        st.session_state["patient_id"] = insert_patient(
                            nome.strip(), telefone.strip(), email.strip(), nascimento, pnotas.strip()
                        )
                        st.success("Paciente criado!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao criar paciente: {e}")
        else:
            idxp = labels.index(sel) - 1
            st.session_state["patient_id"] = patients[idxp]["id"] if idxp >= 0 else None

        st.divider()
        if st.button("🧹 Limpar formulário", use_container_width=True, key=K("att", "clear_form_btn")):
            reset_att_form_state()
            st.success("Formulário limpo.")
            st.rerun()

    patient_id = st.session_state.get("patient_id")
    if not patient_id:
        st.info("Selecione ou crie um paciente na sidebar.")
        st.stop()

    # --- ao trocar de paciente: reseta para não misturar dados e tenta carregar a última anamnese ---
    if st.session_state.get("__att_patient_loaded") != patient_id:
        reset_att_form_state()

        try:
            _latest = list_intakes(patient_id, limit=1)
        except Exception:
            _latest = []
        if _latest:
            try:
                apply_intake_to_form(_latest[0])
                st.session_state["last_intake_id"] = _latest[0].get("id")
            except Exception:
                pass

        st.session_state["__att_patient_loaded"] = patient_id

    # -------------------------
    # Inicializa chaves da V4 (9 origens) se não existirem (compat)
    # -------------------------
    st.session_state.setdefault(K("att", "o9_metodo"), "Aurímetro")
    for o in ORIGENS9:
        oid = o["id"]
        st.session_state.setdefault(K("att", f"o9_{oid}_meas"), 0)
        st.session_state.setdefault(K("att", f"o9_{oid}_note"), "")
        for ev in (o.get("evidence") or []):
            evid = ev.get("id")
            if evid:
                st.session_state.setdefault(K("att", f"o9_{oid}_ev_{evid}"), False)

    # -------------------------
    # Leitura do estado (sem depender de qual aba está aberta)
    # -------------------------
    def _get_sym_0_10() -> Dict[str, int]:
        return {s["id"]: int(st.session_state.get(K("att", s["id"]), 0) or 0) for s in V3_SYMPTOMS}

    def _get_neuro_0_4() -> Dict[str, int]:
        return {q["id"]: int(st.session_state.get(K("att", q["id"]), 0) or 0) for q in V3_NEURO}

    def _get_o9_inputs() -> Tuple[Dict[str, float], Dict[str, Dict[str, bool]], Dict[str, str]]:
        measures: Dict[str, float] = {}
        evid: Dict[str, Dict[str, bool]] = {}
        notes: Dict[str, str] = {}
        for o in ORIGENS9:
            oid = o["id"]
            measures[oid] = float(st.session_state.get(K("att", f"o9_{oid}_meas"), 0) or 0)
            notes[oid] = str(st.session_state.get(K("att", f"o9_{oid}_note"), "") or "")
            emap: Dict[str, bool] = {}
            for ev in (o.get("evidence") or []):
                evid_id = ev.get("id")
                if not evid_id:
                    continue
                emap[evid_id] = bool(st.session_state.get(K("att", f"o9_{oid}_ev_{evid_id}"), False))
            evid[oid] = emap
        return measures, evid, notes

    # anamnese física / contexto (mantém o que já existia, mas em outra organização)
    def _get_phys_meta() -> Dict[str, Any]:
        return {
            "phys_dor_local": str(st.session_state.get(K("att", "phys_dor_local"), "") or ""),
            "phys_dor_score": int(st.session_state.get(K("att", "phys_dor_score"), 0) or 0),
            "phys_dor_regioes": st.session_state.get(K("att", "phys_dor_regioes"), []) or [],
            "phys_hist": str(st.session_state.get(K("att", "phys_hist"), "") or ""),
            "phys_meds_txt": str(st.session_state.get(K("att", "phys_meds_txt"), "") or ""),
            "phys_emocoes_lida": str(st.session_state.get(K("att", "phys_emocoes_lida"), "Prefiro não responder") or "Prefiro não responder"),
            "phys_emocoes_obs": str(st.session_state.get(K("att", "phys_emocoes_obs"), "") or ""),
            "phys_alergias": str(st.session_state.get(K("att", "phys_alergias"), "Não") or "Não"),
            "phys_alergias_quais": str(st.session_state.get(K("att", "phys_alergias_quais"), "") or ""),
            "phys_cirurgias": str(st.session_state.get(K("att", "phys_cirurgias"), "Não") or "Não"),
            "phys_cirurgias_quais": str(st.session_state.get(K("att", "phys_cirurgias_quais"), "") or ""),
            "phys_hist_familia": str(st.session_state.get(K("att", "phys_hist_familia"), "") or ""),
            "phys_conflito_nivel": str(st.session_state.get(K("att", "phys_conflito_nivel"), "Não") or "Não"),
            "phys_conflito_desc": str(st.session_state.get(K("att", "phys_conflito_desc"), "") or ""),
            "phys_transt_alim": str(st.session_state.get(K("att", "phys_transt_alim"), "Não") or "Não"),
            "phys_transt_alim_desc": str(st.session_state.get(K("att", "phys_transt_alim_desc"), "") or ""),
        }

    def _get_flags() -> Dict[str, bool]:
        out: Dict[str, bool] = {}
        for f in FLAGS:
            out[f["id"]] = bool(st.session_state.get(K("att", f["id"]), False))
        return out

    sym = _get_sym_0_10()
    neuro = _get_neuro_0_4()
    time_min = int(st.session_state.get(K("att", "np_time_min"), 5) or 5)
    phys_meta = _get_phys_meta()
    flags = _get_flags()

    measures, evidences, o9_notes = _get_o9_inputs()
    o9_res = compute_origens9(measures, evidences, sym)
    scores = compute_domain_scores_prof(o9_res.get("prob", {}), sym, phys_meta)
    focus = pick_focus(scores, top_n=3)
    qty, cadence = sessions_from_scores(scores)
    readiness_pct = compute_readiness_pct(neuro, time_min)
    qty, cadence, neuro_notes = apply_neuro_to_sessions(qty, cadence, readiness_pct, int(sym.get("sym_dor", 0)))
    pillars = compute_core_pillars(sym)

    # Protocolos / plano (prévia)
    try:
        protocols = load_protocols()
    except Exception as e:
        protocols = {}
        st.warning(f"Não consegui ler protocol_library: {e}")
    selected_names = select_protocols(scores, protocols) if protocols else []
    plan = merge_plan(selected_names, protocols) if protocols else {"chakras_prioritarios": [], "emocoes_prioritarias": [], "cristais_sugeridos": [], "fito_sugerida": [], "alertas": []}

    # Alertas (reaproveita sua lógica existente)
    plan.setdefault("alertas", [])
    def _add_alert(msg: str):
        if msg and msg not in plan["alertas"]:
            plan["alertas"].append(msg)

    if flags.get("flag_back"):
        _add_alert("Dificuldade para deitar de costas: ajuste posição/apoios na cama de cristal.")
    if flags.get("flag_perfume"):
        _add_alert("Sensibilidade a cheiros/perfumes: evite aromas fortes; use aromaterapia bem suave ou omita.")
    if flags.get("flag_heat"):
        _add_alert("Sensibilidade ao calor: mantenha ambiente fresco e confortável.")
    if flags.get("flag_feet"):
        _add_alert("Sensibilidade nos pés: evite pressão intensa; inicie com toques leves.")

    condutas_alerta = build_alert_actions(plan.get("alertas", []), flags, phys_meta, {})
    plan["condutas_alerta"] = condutas_alerta

    # Binaural atual (mantém a aba Binaural como fonte de verdade)
    audio_block = {
        "binaural": {
            "carrier_hz": float(st.session_state.get(KEY_CARRIER, 220.0)),
            "beat_hz": float(st.session_state.get(KEY_BEAT, 10.0)),
            "duracao_s": int(st.session_state.get(KEY_DUR_S, 120)),
        },
        "bg": {
            "gain": float(st.session_state.get(KEY_BG_GAIN, 0.12)),
            "note": "música de fundo é selecionada no computador (não é salva no banco).",
        },
    }
    extra_freq_codes = st.session_state.get("extra_freq_codes") or []

    # -------------------------
    # A) Histórico (compacto)
    # -------------------------
    with st.expander("📚 Histórico (anamneses e planos)", expanded=False):
        try:
            intakes_hist = list_intakes(patient_id, limit=20)
        except Exception as e:
            intakes_hist = []
            st.warning(f"Não consegui carregar anamneses: {e}")

        if intakes_hist:
            rows = []
            for r in intakes_hist:
                sc = _as_dict(r.get("scores_json"))
                top_dom = sorted(sc.items(), key=lambda x: float(x[1]) if isinstance(x[1], (int, float)) else 0.0, reverse=True)[:2]
                top_dom_s = ", ".join([f"{_DOMAIN_LABEL.get(k,k)}:{float(v):.0f}%" for k, v in top_dom]) if top_dom else ""
                # top origens (se existir)
                ans = _as_dict(r.get("answers_json"))
                o9 = ans.get("origens9") or {}
                top_o = []
                try:
                    if isinstance(o9, dict):
                        for oid, ob in o9.items():
                            if isinstance(ob, dict):
                                top_o.append((oid, float(ob.get("meas") or 0)))
                    top_o = sorted(top_o, key=lambda x: x[1], reverse=True)[:2]
                except Exception:
                    top_o = []
                top_o_s = ", ".join([oid for oid, _v in top_o]) if top_o else ""

                rows.append({
                    "quando": str(r.get("created_at") or "")[:19],
                    "queixa": (r.get("complaint") or ""),
                    "top_domínios": top_dom_s,
                    "top_origens": top_o_s,
                    "id": str(r.get("id") or "")[-6:],
                })
            st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

            opts = []
            for r in intakes_hist:
                when = str(r.get("created_at") or "")[:10]
                cid = str(r.get("id") or "")[-4:]
                comp = (r.get("complaint") or "—")
                comp = comp if len(comp) <= 60 else comp[:57] + "..."
                opts.append(f"{when} • {comp} • {cid}")

            sel_i = st.selectbox("Carregar anamnese", opts, key=K("hist", "intake_sel"))
            rsel = intakes_hist[opts.index(sel_i)]
            c1, c2 = st.columns(2)
            if c1.button("Carregar no formulário", type="primary", use_container_width=True, key=K("hist", "load_intake")):
                apply_intake_to_form(rsel)
                st.success("Anamnese carregada.")
                st.rerun()
            if c2.button("Limpar (não apaga histórico)", use_container_width=True, key=K("hist", "clear_form")):
                reset_att_form_state()
                st.success("Formulário limpo.")
                st.rerun()
        else:
            st.caption("Sem anamneses registradas ainda.")

        st.divider()
        st.markdown("**Últimos planos**")
        try:
            plans_hist = list_plans(patient_id, limit=10)
        except Exception as e:
            plans_hist = []
            st.warning(f"Não consegui carregar planos: {e}")

        if plans_hist:
            p0 = plans_hist[0]
            st.write(
                f"Último plano: {str(p0.get('created_at') or '')[:19]} • sessões={p0.get('sessions_qty')} • cadência={p0.get('cadence_days')} dias"
            )
            try:
                sess = list_sessions_nova(p0.get("id"), limit=50)
            except Exception:
                sess = []
            if sess:
                st.dataframe(
                    pd.DataFrame([{"n": r.get("session_n"), "data": r.get("scheduled_date"), "status": r.get("status")} for r in sess]),
                    use_container_width=True,
                    hide_index=True,
                )
        else:
            st.caption("Nenhum plano salvo ainda.")

    # -------------------------
    # B) Fluxo principal (4 abas)
    # -------------------------
    t1, t2, t3, t4 = st.tabs(["🧭 Causas (9 origens)", "📝 Anamnese & contexto", "🧾 Plano & sessões", "🖨️ Impressão"])

    with t1:
        col1, col2 = st.columns([2, 1])
        with col1:
            complaint = st.text_input("Queixa principal (curta)", key=K("att", "complaint"))
        with col2:
            atend_date = st.date_input("Data", value=date.today(), key=K("att", "date"))

        st.info("Preencha **Medição (0–10)** + **Evidências**. O sistema calcula **probabilidade** e **confiança** por origem.")
        mcol1, mcol2 = st.columns([1, 2])
        with mcol1:
            metodo = st.selectbox("Método da medição", ["Aurímetro", "Pêndulo", "Entrevista", "Outro"], key=K("att", "o9_metodo"))
        with mcol2:
            st.caption("Dica: use a nota de cada origem para registrar **o porquê** da medição (ex.: leitura do aurímetro, frase-chave do paciente).")

        cols = st.columns(3)
        for i, o in enumerate(ORIGENS9):
            oid = o["id"]
            with cols[i % 3]:
                with st.expander(f"{i+1}. {o['label']}", expanded=(i < 3)):
                    st.caption(o.get("desc") or "")
                    st.slider("Medição (0–10)", 0, 10, key=K("att", f"o9_{oid}_meas"), help="0 = não aparece · 10 = muito forte")
                    st.markdown("**Evidências**")
                    for ev in (o.get("evidence") or []):
                        evid = ev.get("id")
                        if not evid:
                            continue
                        st.checkbox(ev.get("label") or evid, key=K("att", f"o9_{oid}_ev_{evid}"))
                    st.text_area("Observação (opcional)", height=68, key=K("att", f"o9_{oid}_note"))

        st.divider()
        st.markdown("### Resultado da leitura (automático)")

        # tabela completa
        rows = []
        for o in ORIGENS9:
            oid = o["id"]
            prob = float(o9_res.get("prob", {}).get(oid, 0.0))
            raw = float(o9_res.get("raw", {}).get(oid, 0.0))
            conf = float(o9_res.get("confidence", {}).get(oid, 0.0))
            rows.append({
                "origem": o["label"],
                "score_0_100": round(raw, 1),
                "prob_%": round(prob * 100.0, 1),
                "confiança": f"{_conf_label(conf)} ({conf:.2f})",
                "evidências_marcadas": len(o9_res.get("hits", {}).get(oid, []) or []),
            })
        df_o9 = pd.DataFrame(rows).sort_values(["prob_%", "score_0_100"], ascending=False)
        st.dataframe(df_o9, use_container_width=True, hide_index=True)

        # top 3 + evidências
        st.markdown("### Top 3 origens (com evidências)")
        top3 = df_o9.head(3).to_dict("records") if not df_o9.empty else []
        for r in top3:
            st.markdown(f"**{r['origem']}** — prob **{r['prob_%']}%** · confiança **{r['confiança']}**")
            # progress (prob)
            st.progress(min(1.0, float(r["prob_%"]) / 100.0))
            # mostra evidências
            oid = None
            for o in ORIGENS9:
                if o["label"] == r["origem"]:
                    oid = o["id"]
                    break
            hits = o9_res.get("hits", {}).get(oid, []) if oid else []
            if hits:
                st.caption("Evidências marcadas: " + "; ".join(hits[:6]) + ("…" if len(hits) > 6 else ""))
            else:
                st.caption("Sem evidências marcadas (confiança tende a cair).")

    with t2:
        st.markdown("### Mapa rápido (0–10)")
        st.caption(SCALE_0_10_HELP)

        t_corpo, t_mente, t_pert = st.tabs(["🧍 Corpo", "🧠 Mente", "🫂 Pertencimento"])
        with t_corpo:
            c1, c2 = st.columns(2)
            with c1:
                st.slider("Dor no corpo (geral)", 0, 10, key=K("att", "sym_dor"), help=SCALE_0_10_HELP)
                st.slider("Sono ruim (insônia / acorda cansada)", 0, 10, key=K("att", "sym_sono"), help=SCALE_0_10_HELP)
                st.slider("Sobrecarga / responsabilidades", 0, 10, key=K("att", "sym_sobrecarga"), help=SCALE_0_10_HELP)
            with c2:
                st.slider("Rigidez / travamento", 0, 10, key=K("att", "sym_rigidez"), help=SCALE_0_10_HELP)
                st.slider("Cansaço / fadiga", 0, 10, key=K("att", "sym_fadiga"), help=SCALE_0_10_HELP)

        with t_mente:
            m1, m2 = st.columns(2)
            with m1:
                st.slider("Ansiedade / agitação", 0, 10, key=K("att", "sym_ansiedade"), help=SCALE_0_10_HELP)
                st.slider("Irritabilidade / impaciência", 0, 10, key=K("att", "sym_irritabilidade"), help=SCALE_0_10_HELP)
            with m2:
                st.slider("Pensamentos repetitivos / ruminação", 0, 10, key=K("att", "sym_ruminacao"), help=SCALE_0_10_HELP)

        with t_pert:
            p1, p2 = st.columns(2)
            with p1:
                st.slider("Falta de pertencimento / desconexão", 0, 10, key=K("att", "sym_falta_pertenc"), help=SCALE_0_10_HELP)
                st.slider("Autojulgamento / vergonha", 0, 10, key=K("att", "sym_autocritica"), help=SCALE_0_10_HELP)
            with p2:
                st.slider("Isolamento / dificuldade de pedir ajuda", 0, 10, key=K("att", "sym_isolamento"), help=SCALE_0_10_HELP)

        st.divider()
        st.markdown("### Neuroplasticidade (0–4)")
        st.caption("0 = baixo/difícil agora · 4 = alto/fácil agora (exceto alerta, onde 4 = muito alerta).")
        n1, n2 = st.columns(2)
        with n1:
            st.slider("Meu corpo/mente fica em alerta com facilidade", 0, 4, key=K("att", "np_alerta"))
            st.slider("Consigo pausar 30–60s antes de reagir", 0, 4, key=K("att", "np_pausa"))
        with n2:
            st.slider("Tenho pelo menos 1 recurso que me acalma", 0, 4, key=K("att", "np_recurso"))
            st.slider("Consigo praticar um exercício curto todos os dias", 0, 4, key=K("att", "np_adesao"))

        # tempo disponível (min/dia)
        time_labels = [lab for _m, lab in V3_TIME_OPTIONS]
        time_vals = [int(_m) for _m, _lab in V3_TIME_OPTIONS]
        default_min = st.session_state.get(K("att", "np_time_min"), 5)
        if default_min not in time_vals:
            default_min = 5
        idx_def = time_vals.index(default_min)
        sel_lab = st.selectbox("Tempo disponível para prática (em casa)", time_labels, index=idx_def, key=K("att", "np_time_lab"))
        st.session_state[K("att", "np_time_min")] = int(time_vals[time_labels.index(sel_lab)])

        st.divider()
        st.markdown("### Sinais de atenção")
        fcols = st.columns(2)
        for i, f in enumerate(FLAGS):
            with fcols[i % 2]:
                st.checkbox(f["label"], key=K("att", f["id"]))

        with st.expander("🩺 Anamnese física (detalhes)", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                st.text_input("Dor / queixa principal (onde dói?)", key=K("att", "phys_dor_local"))
            with c2:
                st.slider("Intensidade da dor (0=sem dor; 10=máxima)", 0, 10, key=K("att", "phys_dor_score"))
            st.multiselect("Regiões afetadas (marque se fizer sentido)", PHYS_DOR_REGIOES, key=K("att", "phys_dor_regioes"))
            st.text_area("Histórico de saúde / cirurgias relevantes", height=80, key=K("att", "phys_hist"))
            st.text_area("Medicamentos / tratamentos atuais", height=80, key=K("att", "phys_meds_txt"))

            st.markdown("**Aspectos emocionais e contexto**")
            EMO_OPTS = ["Prefiro não responder", "Guardo pra mim / engulo", "Falo / peço ajuda", "Explodo / fico irritada", "Choro / fico retraída", "Processo com terapia/meditação", "Outro"]
            CONFLITO_OPTS = ["Não", "Leve", "Moderado", "Grave"]
            SIMNAO_OPTS = ["Não", "Sim"]
            TRANST_ALIM_OPTS = ["Não", "Suspeita/Em investigação", "Sim"]

            if st.session_state.get(K("att", "phys_emocoes_lida")) not in EMO_OPTS:
                st.session_state[K("att", "phys_emocoes_lida")] = EMO_OPTS[0]
            if st.session_state.get(K("att", "phys_conflito_nivel")) not in CONFLITO_OPTS:
                st.session_state[K("att", "phys_conflito_nivel")] = CONFLITO_OPTS[0]
            if st.session_state.get(K("att", "phys_alergias")) not in SIMNAO_OPTS:
                st.session_state[K("att", "phys_alergias")] = SIMNAO_OPTS[0]
            if st.session_state.get(K("att", "phys_cirurgias")) not in SIMNAO_OPTS:
                st.session_state[K("att", "phys_cirurgias")] = SIMNAO_OPTS[0]
            if st.session_state.get(K("att", "phys_transt_alim")) not in TRANST_ALIM_OPTS:
                st.session_state[K("att", "phys_transt_alim")] = TRANST_ALIM_OPTS[0]

            d1, d2 = st.columns(2)
            with d1:
                st.selectbox("Como você lida com suas emoções?", EMO_OPTS, key=K("att", "phys_emocoes_lida"))
            with d2:
                st.selectbox("Atualmente possui conflito familiar?", CONFLITO_OPTS, key=K("att", "phys_conflito_nivel"))
            st.text_area("Se sim, descreva (opcional)", height=68, key=K("att", "phys_conflito_desc"))
            st.text_input("Observações sobre emoções (opcional)", key=K("att", "phys_emocoes_obs"))

            st.markdown("**Saúde e antecedentes**")
            e1, e2 = st.columns(2)
            with e1:
                st.selectbox("Tem alguma alergia?", SIMNAO_OPTS, key=K("att", "phys_alergias"))
                st.text_input("Se sim, qual(is)?", key=K("att", "phys_alergias_quais"))
            with e2:
                st.selectbox("Já fez alguma cirurgia?", SIMNAO_OPTS, key=K("att", "phys_cirurgias"))
                st.text_input("Se sim, qual(is)?", key=K("att", "phys_cirurgias_quais"))

            st.text_area("Histórico familiar relevante (se houver)", height=80, key=K("att", "phys_hist_familia"))
            f1, f2 = st.columns(2)
            with f1:
                st.selectbox("Possui transtorno alimentar?", TRANST_ALIM_OPTS, key=K("att", "phys_transt_alim"))
            with f2:
                st.text_input("Se sim/suspeita, qual/observações?", key=K("att", "phys_transt_alim_desc"))

        st.text_area("Notas do terapeuta (opcional)", height=100, key=K("att", "notes"))

    with t3:
        st.markdown("### Painel rápido")
        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("Dor/Fibro", f"{pillars['Dor/Fibro']:.0f}%")
        k2.metric("Sistema nervoso", f"{pillars['Sistema nervoso']:.0f}%")
        k3.metric("Burnout", f"{pillars['Burnout/Cansaço']:.0f}%")
        k4.metric("Pertencimento", f"{pillars['Pertencimento']:.0f}%")
        k5.metric("Prontidão", f"{readiness_pct:.0f}%")

        st.divider()
        st.markdown("### Domínios (0–100) • usados para selecionar protocolos")
        df_scores = pd.DataFrame(
            [{"domínio": _DOMAIN_LABEL.get(k, k), "score_%": round(v, 1)} for k, v in sorted(scores.items(), key=lambda x: x[1], reverse=True)]
        )
        st.dataframe(df_scores, use_container_width=True, hide_index=True)

        st.markdown("### Foco (Top 3)")
        if focus:
            st.dataframe(
                pd.DataFrame(
                    [{"prioridade": i + 1, "domínio": _DOMAIN_LABEL.get(d, d), "score_%": round(sc, 1)} for i, (d, sc) in enumerate(focus)]
                ),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.caption("—")

        st.divider()
        st.markdown("### Sessões sugeridas")
        st.write(f"**Quantidade:** {qty}  •  **Cadência:** {cadence} dias")
        if neuro_notes:
            st.caption("Ajustes por neuroplasticidade: " + " · ".join(neuro_notes))

        st.markdown("### Protocolos selecionados")
        if selected_names:
            st.dataframe(pd.DataFrame([{"protocolo": n} for n in selected_names]), use_container_width=True, hide_index=True)
        else:
            st.caption("Sem protocolos (verifique se protocol_library está disponível).")

        st.markdown("### Plano consolidado (resumo)")
        plan_rows = [
            {"categoria": "Chakras prioritários", "itens": _join_list(plan.get("chakras_prioritarios"), sep="; ")},
            {"categoria": "Emoções prioritárias", "itens": _join_list(plan.get("emocoes_prioritarias"), sep="; ")},
            {"categoria": "Cristais sugeridos", "itens": _join_list(plan.get("cristais_sugeridos"), sep="; ")},
            {"categoria": "Fito sugerida", "itens": _join_list(plan.get("fito_sugerida"), sep="; ")},
            {"categoria": "Alertas / cuidados", "itens": _join_list(plan.get("alertas"), sep="; ")},
        ]
        st.dataframe(pd.DataFrame(plan_rows), use_container_width=True, hide_index=True)

        st.markdown("### 🚨 Alertas e condutas")
        if plan.get("condutas_alerta"):
            df_ca = pd.DataFrame(plan.get("condutas_alerta"), columns=["Prioridade", "Categoria", "Detalhe", "Conduta sugerida"])
            st.dataframe(df_ca, use_container_width=True, hide_index=True)
        else:
            st.caption("Sem alertas relevantes (além do básico).")

        st.divider()
        st.markdown("### Botões finais")
        b1, b2 = st.columns(2)

        # answers_store: salva tudo organizado
        answers_store: Dict[str, Any] = {}
        answers_store["o9_metodo"] = str(st.session_state.get(K("att", "o9_metodo"), "Aurímetro") or "Aurímetro")
        answers_store["origens9"] = {
            o["id"]: {
                "meas": int(st.session_state.get(K("att", f"o9_{o['id']}_meas"), 0) or 0),
                "note": str(st.session_state.get(K("att", f"o9_{o['id']}_note"), "") or ""),
                "evidence": {
                    (ev.get("id") or ""): bool(st.session_state.get(K("att", f"o9_{o['id']}_ev_{ev.get('id')}"), False))
                    for ev in (o.get("evidence") or [])
                    if ev.get("id")
                },
            }
            for o in ORIGENS9
        }
        # V3 symptoms + neuro + tempo
        for s in V3_SYMPTOMS:
            answers_store[s["id"]] = int(st.session_state.get(K("att", s["id"]), 0) or 0)
        for q in V3_NEURO:
            answers_store[q["id"]] = int(st.session_state.get(K("att", q["id"]), 0) or 0)
        answers_store["np_time_min"] = int(st.session_state.get(K("att", "np_time_min"), 5) or 5)
        # físico
        answers_store.update(phys_meta or {})

        notes = st.session_state.get(K("att", "notes"), "") or ""
        complaint = st.session_state.get(K("att", "complaint"), "") or ""
        atend_date = st.session_state.get(K("att", "date"), date.today())

        with b1:
            if st.button("Salvar anamnese", use_container_width=True, key=K("att", "save_intake")):
                try:
                    intake_id = insert_intake(patient_id, complaint, answers_store, scores, flags, notes)
                    st.session_state["last_intake_id"] = intake_id
                    st.success("Anamnese salva!")
                except Exception as e:
                    st.error(f"Erro ao salvar anamnese: {e}")

        with b2:
            if st.button("Gerar plano terapêutico + criar sessões", type="primary", use_container_width=True, key=K("att", "gen_plan")):
                try:
                    intake_id = st.session_state.get("last_intake_id")
                    if not intake_id:
                        intake_id = insert_intake(patient_id, complaint, answers_store, scores, flags, notes)
                        st.session_state["last_intake_id"] = intake_id

                    scripts = build_session_scripts(qty, cadence, focus, selected_names, protocols, audio_block, extra_freq_codes)
                    for s in scripts:
                        s["alertas"] = plan.get("alertas", [])
                        s["condutas_alerta"] = plan.get("condutas_alerta", [])
                        s["phys_meta"] = phys_meta
                        s["origens9_resultado"] = o9_res

                    plan_id = insert_plan(
                        patient_id=patient_id,
                        intake_id=intake_id,
                        focus=focus,
                        selected_names=selected_names,
                        sessions_qty=qty,
                        cadence_days=cadence,
                        plan_json={
                            "date": str(atend_date),
                            "complaint": complaint,
                            "scores": scores,
                            "focus": focus,
                            "origens9_resultado": o9_res,
                            "pillars_v3": pillars,
                            "readiness_pct": readiness_pct,
                            "np_time_min": time_min,
                            "answers": answers_store,
                            "selected_protocols": selected_names,
                            "plan": plan,
                            "audio": audio_block,
                            "frequencias": [{"code": c} for c in extra_freq_codes],
                        },
                    )
                    for s in scripts:
                        insert_session_nova(plan_id, patient_id, int(s["session_n"]), s["scheduled_date"], s["status"], s)

                    st.success(f"Plano criado e sessões geradas em sessions_nova! plan_id={plan_id}")
                    st.rerun()
                except Exception as e:
                    st.error(f"Erro ao gerar plano/sessões: {e}")

    with t4:
        st.subheader("🖨️ Receituário para impressão")
        with st.expander("Gerar receituário (puxa as informações salvas do paciente)", expanded=True):
            try:
                _plans = list_plans(patient_id)
            except Exception as e:
                _plans = []
                st.warning(f"Não consegui listar planos: {e}")

            if not _plans:
                st.info("Ainda não há plano salvo para este paciente. Gere um plano e salve antes de imprimir.")
            else:
                def _plan_label(p):
                    dt = p.get("created_at") or (p.get("plan_json") or {}).get("date") or ""
                    dt = _fmt_date_br(dt) if dt else ""
                    pid = str(p.get("id") or "")[-6:]
                    qty2 = p.get("sessions_qty") or ""
                    cad2 = p.get("cadence_days") or ""
                    extra = f" • {qty2} sessões/{cad2}d" if qty2 and cad2 else ""
                    return f"{dt or 'sem data'} — Plano {pid}{extra}"

                plan_labels = [_plan_label(p) for p in _plans]
                rx_sel = st.selectbox("Plano para imprimir", plan_labels, index=0, key=K("rx", "plan_sel"))
                plan_row = _plans[plan_labels.index(rx_sel)]
                plan_id = plan_row.get("id")

                try:
                    sess_rows = list_sessions_nova(plan_id)
                except Exception as e:
                    sess_rows = []
                    st.warning(f"Não consegui listar sessions_nova: {e}")

                pat = get_patient(patient_id) or {"id": patient_id}
                rx_data = _build_receituario_data_from_plan(pat, plan_row, sess_rows)

                st.caption("Dica: mantenha o template DOCX no mesmo diretório do app (Receituario_Claudiafito_Template.docx) ou envie abaixo.")
                tpl_up = st.file_uploader("Template DOCX (opcional)", type=["docx"], key=K("rx", "tpl"))

                colrx1, colrx2 = st.columns(2)
                with colrx1:
                    if st.button("Gerar receituário (DOCX)", use_container_width=True, key=K("rx", "gen_docx")):
                        try:
                            tpl_io = io.BytesIO(tpl_up.read()) if tpl_up else None
                            st.session_state["rx_docx_bytes"] = generate_receituario_docx_bytes(rx_data, template_file=tpl_io)
                            st.success("DOCX gerado.")
                        except Exception as e:
                            st.error(f"Erro ao gerar DOCX: {e}")

                with colrx2:
                    if st.button("Gerar receituário (PDF)", use_container_width=True, key=K("rx", "gen_pdf")):
                        try:
                            st.session_state["rx_pdf_bytes"] = generate_receituario_pdf_bytes(rx_data)
                            st.success("PDF gerado.")
                        except Exception as e:
                            st.error(f"Erro ao gerar PDF: {e}")

                dcol1, dcol2 = st.columns(2)
                with dcol1:
                    if st.session_state.get("rx_docx_bytes"):
                        st.download_button(
                            "⬇️ Baixar DOCX preenchido",
                            data=st.session_state["rx_docx_bytes"],
                            file_name=f"receituario_{(pat.get('nome') or 'paciente').strip().replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=K("rx", "dl_docx"),
                        )
                with dcol2:
                    if HAS_REPORTLAB and st.session_state.get("rx_pdf_bytes"):
                        st.download_button(
                            "⬇️ Baixar PDF",
                            data=st.session_state["rx_pdf_bytes"],
                            file_name=f"receituario_{(pat.get('nome') or 'paciente').strip().replace(' ','_')}.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                            key=K("rx", "dl_pdf"),
                        )

                with st.expander("Prévia do que vai no receituário", expanded=False):
                    st.write("Paciente:", rx_data.get("patient_nome"))
                    st.write("Queixa:", rx_data.get("queixa"))
                    st.write("Foco:", rx_data.get("focus"))
                    st.write("Binaural:", rx_data.get("binaural_txt"))
                    st.write("Frequências auxiliares:", rx_data.get("freq_aux_txt"))
                    st.write("Cama de cristal:", rx_data.get("cama_txt"))
                    st.write("Cristais:", rx_data.get("cristais_txt"))
                    st.write("Fito:", rx_data.get("fito_txt"))
